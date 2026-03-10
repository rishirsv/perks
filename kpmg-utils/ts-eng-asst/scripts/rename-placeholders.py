"""
Rename Placeholders in Engagement Letter Templates

Uses reference/placeholder-mapping.json to batch-rename all placeholders
in both .docx files to UPPER_SNAKE_CASE. Handles context-dependent
disambiguation for {{Date}}, {{Name}}, {{xxx}}, and {{xx}} placeholders.

Usage:
    python3 scripts/rename-placeholders.py
"""

import json
import re
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document

PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")

# Unicode smart character normalization for matching
_SMART_MAP = {
    "\u201c": '"',   # left double quote
    "\u201d": '"',   # right double quote
    "\u2018": "'",   # left single quote
    "\u2019": "'",   # right single quote
    "\u2013": "-",   # en dash
    "\u2014": "-",   # em dash
    "\u00a0": " ",   # non-breaking space
}

def _normalize_text(text: str) -> str:
    """Normalize Unicode smart chars to ASCII for matching."""
    for smart, plain in _SMART_MAP.items():
        text = text.replace(smart, plain)
    return text


def load_mapping() -> dict:
    mapping_path = PROJECT_ROOT / "reference" / "placeholder-mapping.json"
    with open(mapping_path, "r", encoding="utf-8") as f:
        return json.load(f)


def _replace_in_runs(runs, old: str, new: str) -> int:
    """Replace old placeholder with new in runs. Returns count of replacements."""
    count = 0
    for run in runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    if count > 0:
        return count

    # Handle placeholders split across runs
    full_text = "".join(run.text for run in runs)
    if old not in full_text:
        return 0

    while True:
        full_text = "".join(run.text for run in runs)
        idx = full_text.find(old)
        if idx == -1:
            break
        end_idx = idx + len(old)
        cursor = 0
        inserted = False

        for run in runs:
            text = run.text
            run_end = cursor + len(text)

            if run_end <= idx or cursor >= end_idx:
                cursor = run_end
                continue

            before = text[: max(0, idx - cursor)] if idx > cursor else ""
            after = text[end_idx - cursor:] if run_end > end_idx else ""

            if not inserted:
                run.text = before + new + after
                inserted = True
            else:
                run.text = after

            cursor = run_end

        count += 1

    return count


def _replace_first_in_runs(runs, old: str, new: str) -> bool:
    """Replace only the FIRST occurrence of old with new across runs."""
    full_text = "".join(run.text for run in runs)
    idx = full_text.find(old)
    if idx == -1:
        return False

    end_idx = idx + len(old)
    cursor = 0
    inserted = False

    for run in runs:
        text = run.text
        run_end = cursor + len(text)

        if run_end <= idx or cursor >= end_idx:
            cursor = run_end
            continue

        before = text[: max(0, idx - cursor)] if idx > cursor else ""
        after = text[end_idx - cursor:] if run_end > end_idx else ""

        if not inserted:
            run.text = before + new + after
            inserted = True
        else:
            run.text = after

        cursor = run_end

    return True


def apply_simple_renames(doc: Document, mappings: list[dict], template_type: str) -> dict:
    """Apply simple (non-context-dependent) renames. Returns counts."""
    stats = {"replaced": 0, "skipped": 0}

    # Build lookup: old -> new, filtering by applies_to
    rename_map = {}
    for m in mappings:
        applies = m.get("applies_to", "both")
        if applies == "both" or applies == template_type:
            old = m["old"]
            new = m["new"]
            # Skip context-dependent ones (handled separately)
            if old in ("{{Date}}", "{{DATE}}", "{{Name}}", "{{xxx}}", "{{xx}}", "{{Partner}}", "{{Senior Manager}}"):
                continue
            rename_map[old] = new

    # Body paragraphs
    for para in doc.paragraphs:
        if not para.runs:
            continue
        for old, new in rename_map.items():
            count = _replace_in_runs(para.runs, old, new)
            stats["replaced"] += count

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.runs:
                        continue
                    for old, new in rename_map.items():
                        count = _replace_in_runs(para.runs, old, new)
                        stats["replaced"] += count

    # Headers/footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            if not para.runs:
                continue
            for old, new in rename_map.items():
                count = _replace_in_runs(para.runs, old, new)
                stats["replaced"] += count
        for para in section.footer.paragraphs:
            if not para.runs:
                continue
            for old, new in rename_map.items():
                count = _replace_in_runs(para.runs, old, new)
                stats["replaced"] += count

    return stats


def apply_context_renames(doc: Document, context_rules: list[dict]) -> int:
    """Apply context-dependent renames using paragraph index and context text."""
    count = 0
    paras = list(doc.paragraphs)

    for rule in context_rules:
        p_idx = rule.get("paragraph_index")
        ctx = rule.get("context_contains")
        old = rule["old"]
        new = rule["new"]
        position = rule.get("position")

        if p_idx is not None and p_idx < len(paras):
            para = paras[p_idx]
            if not para.runs:
                continue

            # Check context matches (if specified)
            if ctx and ctx not in para.text:
                # Try adjacent paragraphs (index might shift slightly)
                found = False
                for offset in [-1, 1, -2, 2]:
                    adj_idx = p_idx + offset
                    if 0 <= adj_idx < len(paras):
                        adj = paras[adj_idx]
                        if old in adj.text and (ctx is None or ctx in adj.text):
                            para = adj
                            found = True
                            break
                if not found and ctx:
                    # Search all paragraphs for context match
                    for i, p in enumerate(paras):
                        if old in p.text and ctx in p.text:
                            para = p
                            break

            if position == "second":
                # Replace first occurrence with a temp marker, then replace second, then restore
                temp = "___TEMP_FIRST___"
                if _replace_first_in_runs(para.runs, old, temp):
                    if _replace_first_in_runs(para.runs, old, new):
                        count += 1
                    # Restore first occurrence
                    _replace_in_runs(para.runs, temp, old)
            elif position == "first":
                if _replace_first_in_runs(para.runs, old, new):
                    count += 1
            else:
                c = _replace_in_runs(para.runs, old, new)
                count += c

    return count


def apply_smart_char_sweep(doc: Document, mappings: list[dict], template_type: str) -> int:
    """
    Final sweep: for any placeholder that wasn't matched via simple renames
    (due to smart quotes/apostrophes/dashes in .docx), find by normalized text
    and replace the actual runs.
    """
    count = 0
    # Build a normalized lookup
    norm_map = {}
    for m in mappings:
        applies = m.get("applies_to", "both")
        if applies == "both" or applies == template_type:
            old_norm = _normalize_text(m["old"])
            norm_map[old_norm] = m["new"]

    def _sweep_paragraph(para):
        nonlocal count
        if not para.runs:
            return
        full_text = "".join(r.text for r in para.runs)
        # Find all {{...}} placeholders that are NOT yet UPPER_SNAKE_CASE
        for match in PLACEHOLDER_RE.finditer(full_text):
            ph = match.group()
            if re.match(r"\{\{[A-Z_0-9]+\}\}$", ph):
                continue  # already renamed
            # Try normalized match
            ph_norm = _normalize_text(ph)
            if ph_norm in norm_map:
                new = norm_map[ph_norm]
                c = _replace_in_runs(para.runs, ph, new)
                count += c

    for para in doc.paragraphs:
        _sweep_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _sweep_paragraph(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _sweep_paragraph(para)
        for para in section.footer.paragraphs:
            _sweep_paragraph(para)

    return count


def rename_template(template_type: str, mapping: dict) -> None:
    """Rename all placeholders in one template."""
    if template_type == "buyside":
        path = PROJECT_ROOT / "dist" / "buyside-engagement-letter.docx"
    else:
        path = PROJECT_ROOT / "dist" / "sellside-engagement-letter.docx"

    print(f"\n{'=' * 60}")
    print(f"RENAMING: {template_type} ({path.name})")
    print(f"{'=' * 60}")

    doc = Document(str(path))

    # Step 1: Apply simple (non-ambiguous) renames
    simple_stats = apply_simple_renames(doc, mapping["mappings"], template_type)
    print(f"  Simple renames: {simple_stats['replaced']} replacements")

    # Step 2: Apply context-dependent renames
    ctx_dep = mapping.get("context_dependent_renames", {})
    ctx_count = 0

    if template_type == "buyside":
        ctx_groups = ["buyside_date_renames", "buyside_name_renames", "buyside_xxx_renames", "buyside_xx_renames"]
    else:
        ctx_groups = ["sellside_date_renames", "sellside_xxx_renames", "sellside_date_deliver"]

    for group_key in ctx_groups:
        rules = ctx_dep.get(group_key, [])
        c = apply_context_renames(doc, rules)
        ctx_count += c
        print(f"  Context renames ({group_key}): {c} replacements")

    # Step 2b: Handle simple sellside-specific fields that share names with buyside
    if template_type == "sellside":
        paras = list(doc.paragraphs)
        # {{Partner}} at "Lead Engagement Partner" line
        for p in paras:
            if "Lead Engagement Partner" in p.text and "{{Partner}}" in p.text:
                c = _replace_in_runs(p.runs, "{{Partner}}", "{{TEAM_PARTNER_NAME}}")
                ctx_count += c
        # {{Senior Manager}} at "Engagement Senior Manager" line
        for p in paras:
            if "Engagement Senior Manager" in p.text and "{{Senior Manager}}" in p.text:
                c = _replace_in_runs(p.runs, "{{Senior Manager}}", "{{TEAM_SENIOR_MANAGER_NAME}}")
                ctx_count += c
        # Remaining {{Date}} in header
        for section in doc.sections:
            for p in section.header.paragraphs:
                c = _replace_in_runs(p.runs, "{{Date}}", "{{LETTER_DATE}}")
                ctx_count += c
        # Body standalone {{Date}} at para 7
        if len(paras) > 7:
            p = paras[7]
            if p.text.strip() == "{{Date}}":
                c = _replace_in_runs(p.runs, "{{Date}}", "{{LETTER_DATE}}")
                ctx_count += c
        # Remaining {{Date}} in late sections of sellside
        for p in paras:
            if "{{Date}}" in p.text:
                c = _replace_in_runs(p.runs, "{{Date}}", "{{LETTER_DATE}}")
                ctx_count += c

    if template_type == "buyside":
        # Handle header Date
        for section in doc.sections:
            for p in section.header.paragraphs:
                c = _replace_in_runs(p.runs, "{{Date}}", "{{LETTER_DATE}}")
                ctx_count += c
        # Handle remaining {{xxx}} (FEE_FDD_HIGH) and {{xx}} (FEE_TAX_HIGH)
        paras = list(doc.paragraphs)
        for p in paras:
            if "{{xxx}}" in p.text and "Financial due diligence" in p.text:
                c = _replace_in_runs(p.runs, "{{xxx}}", "{{FEE_FDD_HIGH}}")
                ctx_count += c
            if "{{xx}}" in p.text and "Tax due diligence" in p.text:
                c = _replace_in_runs(p.runs, "{{xx}}", "{{FEE_TAX_HIGH}}")
                ctx_count += c

    if template_type == "sellside":
        # Handle remaining {{xx}} in fee discount clause
        paras = list(doc.paragraphs)
        for p in paras:
            if "{{xx}}" in p.text:
                c = _replace_in_runs(p.runs, "{{xx}}", "{{FEE_DISCOUNT_PERCENT}}")
                ctx_count += c

    print(f"  Total context renames: {ctx_count}")

    # Step 2c: Smart character sweep for placeholders with Unicode quotes/dashes
    sweep_count = apply_smart_char_sweep(doc, mapping["mappings"], template_type)
    print(f"  Smart char sweep: {sweep_count} replacements")

    # Step 2d: Handle nested/malformed placeholders (placeholder inside placeholder)
    nested_count = 0
    paras = list(doc.paragraphs)
    if template_type == "buyside":
        # {{our due diligence report(s)/deliverables relating to Project {{PROJECT_CODE_NAME}}
        # This is a nested open-brace situation. Replace the whole thing.
        for p in paras:
            full = "".join(r.text for r in p.runs) if p.runs else p.text
            target = "{{our due diligence report(s)/deliverables relating to Project {{PROJECT_CODE_NAME}}"
            norm_target = _normalize_text(target)
            norm_full = _normalize_text(full)
            if norm_target in norm_full:
                # Find the actual text in runs
                actual_full = "".join(r.text for r in p.runs)
                # Find the start of the nested placeholder
                idx = _normalize_text(actual_full).find(norm_target)
                if idx >= 0:
                    actual_target = actual_full[idx:idx + len(target)]
                    c = _replace_in_runs(p.runs, actual_target, "{{RELEASE_DELIVERABLES_DESC}}")
                    nested_count += c
    if template_type == "sellside":
        # {{although as of the date of this letter, KPMG has not been engaged by Bidder with respect to Project {{PROJECT_CODE_NAME}}
        for p in paras:
            full = "".join(r.text for r in p.runs) if p.runs else p.text
            target = "{{although as of the date of this letter, KPMG has not been engaged by Bidder with respect to Project {{PROJECT_CODE_NAME}}"
            norm_target = _normalize_text(target)
            norm_full = _normalize_text(full)
            if norm_target in norm_full:
                actual_full = "".join(r.text for r in p.runs)
                idx = _normalize_text(actual_full).find(norm_target)
                if idx >= 0:
                    actual_target = actual_full[idx:idx + len(target)]
                    c = _replace_in_runs(p.runs, actual_target, "{{EXISTING_SERVICES_CAVEAT}}")
                    nested_count += c
    print(f"  Nested placeholder fixes: {nested_count} replacements")

    # Step 3: Validate — check remaining placeholders
    remaining = []
    for i, para in enumerate(doc.paragraphs):
        for match in PLACEHOLDER_RE.finditer(para.text):
            remaining.append((i, match.group(), para.text.strip()[:80]))
    for section in doc.sections:
        for para in section.header.paragraphs:
            for match in PLACEHOLDER_RE.finditer(para.text):
                remaining.append(("header", match.group(), para.text.strip()[:80]))
        for para in section.footer.paragraphs:
            for match in PLACEHOLDER_RE.finditer(para.text):
                remaining.append(("footer", match.group(), para.text.strip()[:80]))

    # Check how many are now UPPER_SNAKE_CASE
    renamed = [r for r in remaining if re.match(r"\{\{[A-Z_0-9]+\}\}", r[1])]
    unrenamed = [r for r in remaining if not re.match(r"\{\{[A-Z_0-9]+\}\}", r[1])]

    print(f"\n  Remaining placeholders: {len(remaining)} total")
    print(f"    Renamed (UPPER_SNAKE): {len(renamed)}")
    print(f"    NOT yet renamed:       {len(unrenamed)}")

    if unrenamed:
        print(f"\n  *** Unrenamed placeholders ***")
        for loc, ph, ctx in unrenamed:
            print(f"    [¶{loc}] {ph}")
            print(f"           {ctx}")

    # Save
    doc.save(str(path))
    print(f"\n  Saved: {path}")


def main():
    mapping = load_mapping()
    rename_template("buyside", mapping)
    rename_template("sellside", mapping)
    print(f"\n{'=' * 60}")
    print("Done! Run scripts/audit-placeholders.py to verify.")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
