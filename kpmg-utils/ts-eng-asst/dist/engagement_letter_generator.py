"""
Engagement Letter Generator (self-contained)

Self-contained script for ChatGPT Code Interpreter. Generates a completed
engagement letter by:
  1. Loading a .docx template
  2. Removing guidance blocks ({{GUIDANCE_NN}} paragraphs)
  3. Replacing the FDD scope block with industry-specific content
  4. Replacing all field/choice placeholders
  5. Validating no {{...}} tokens remain
  6. Saving the filled .docx

Usage (Code Interpreter):
    variables = {"CLIENT_LEGAL_NAME": "Acme Inc.", ...}
    scope_selection = {
        # Optional: omit unchecked top-level scope items (default: include all)
        "excluded_top_level_ids": [],          # e.g. ["scope.002"]
        # Optional (preferred for section-level UX): omit complete scope sections
        "excluded_section_keys": [],           # e.g. ["net_debt", "working_capital"]
    }
    result = generate_engagement_letter(
        template_file="buyside-engagement-letter.docx",
        scope_library_file="scope-library.json",
        industry="healthcare",
        variables=variables,
        scope_selection=scope_selection,
        output_file="Engagement_Letter_Acme_Inc.docx"
    )

Usage (CLI):
    python3 engagement_letter_generator.py \\
        --template buyside-engagement-letter.docx \\
        --scope-library scope-library.json \\
        --industry healthcare \\
        --variables '{"CLIENT_LEGAL_NAME": "Acme Inc.", ...}' \\
        --scope-selection '{"excluded_top_level_ids":["scope.002"]}' \\
        --output output.docx
"""

import json
import difflib
import re
import sys
from datetime import date
from dataclasses import dataclass, field
from copy import deepcopy
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.text.paragraph import Paragraph
from scope_engine import (
    apply_optional_scope_modules as _core_apply_optional_scope_modules,
    build_industry_scope_view as _core_build_industry_scope_view,
    ordered_active_section_keys as _core_ordered_active_section_keys,
    parse_scope_selection as _core_parse_scope_selection,
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")
GUIDANCE_RE = re.compile(r"\{\{GUIDANCE_\d+\}\}")
GUIDANCE_FREEFORM_RE = re.compile(r"(?i)\{\{\s*GUIDANCE:")

# Placeholders that must be explicitly provided by the user (no inference).
# If present in the selected template and missing/blank at generation time, we refuse.
MUST_ASK_KEYS: set[str] = {
    "CHOICE_DEAL_TYPE",  # Buyside-only; enforced only when present in template.
}

# Indentation (EMU) matching engagement letter templates
_INDENT_BULLET = Emu(457200)     # 0.50 in — top-level bullet
_INDENT_SUBBULLET = Emu(685800)  # 0.75 in — sub-level bullet
_INDENT_SUBSUBBULLET = Emu(914400)  # 1.00 in — sub-sub-level bullet (rare)

# All guidance placeholder keys (deleted, not replaced)
GUIDANCE_KEYS = {
    f"{{{{GUIDANCE_{str(i).zfill(2)}}}}}" for i in range(1, 19)
}

_REVIEW_PREFIX = "[REVIEW:"

# Block markers (standalone paragraphs) used for deterministic section removal.
BLOCK_INDEPENDENCE_START = "{{BLOCK_INDEPENDENCE_START}}"
BLOCK_INDEPENDENCE_END = "{{BLOCK_INDEPENDENCE_END}}"

# Placeholders where an explicit empty string is a valid default (omits a clause).
ALLOW_EMPTY_KEYS = {
    "EXISTING_SERVICES_CAVEAT",
}

# Placeholder keys that should remain untouched in the template output.
# These are release-letter fields completed outside this generation flow.
PRESERVE_TEMPLATE_PLACEHOLDERS = {
    "CHOICE_DEAL_TYPE_2",
    "CHOICE_DEAL_TYPE_3",
    "RELEASE_DATE",
    "THIRD_PARTY_LEGAL_NAME",
    "RECIPIENT_LEGAL_NAME",
    "RECIPIENT_SIGNATORY_NAME",
    "RELEASE_DELIVERABLES_DESC",
    "SIGNING_PARTNER_NAME_2",
}

# Numeric money-like fields where we normalize formatting (commas) and support
# legacy "in thousands" inputs (e.g., "80" -> "80,000").
MONEY_KEYS = {
    "FEE_FDD_LOW",
    "FEE_FDD_HIGH",
    "FEE_TAX_LOW",
    "FEE_TAX_HIGH",
}

_MONEY_INT_RE = re.compile(r"^\s*\$?\s*([0-9][0-9,]*)\s*$")

# Common user-facing industry aliases/variants mapped to canonical keys.
_INDUSTRY_ALIAS_MAP = {
    "telecom": "telecomm",
    "telecommunications": "telecomm",
    "professional services": "prof_services",
    "professional_service": "prof_services",
    "professional_services": "prof_services",
    "prof services": "prof_services",
    "pro services": "prof_services",
    "real estate": "real_estate",
    "realestate": "real_estate",
    "eye care": "eyecare",
}

_HEADING_ACRONYMS = {
    "arr": "ARR",
    "spa": "SPA",
    "gaap": "GAAP",
}


_DEFAULT_CONFIDENTIALITY_PROCEDURES = (
    "• The KPMG engagement team assisting Client will exclude any individual who is currently providing "
    "the outlined services to Target;\n"
    "• All files related to the Transaction, including computer files, of the Client engagement team will "
    "be segregated from KPMG's regular files and will be accessible only by members of the Client "
    "engagement team;\n"
    "• Members of the Client engagement team will be physically separate (and working virtually) from all "
    "individuals providing services to Target; and\n"
    "• All members of the engagement team assisting Client have been and any individual that joins the "
    "engagement team, will be, informed of the above procedures and have agreed to abide by such procedures."
)


def _format_heading_from_key(value: str) -> str:
    key = str(value or "").strip().lower()
    parts = [p for p in key.replace("-", "_").split("_") if p]
    if not parts:
        return str(value or "").replace("_", " ").strip().title()
    return " ".join(_HEADING_ACRONYMS.get(p, p.capitalize()) for p in parts)


def _is_blank(value) -> bool:
    return value is None or str(value).strip() == ""


def _set_if_blank(out: dict, key: str, value) -> None:
    if _is_blank(out.get(key)):
        out[key] = value


def _format_today(d: Optional[date] = None) -> str:
    d = d or date.today()
    # Platform-independent Month D, YYYY (no leading zero on day).
    return f"{d:%B} {d.day}, {d.year}"


_LEGAL_SUFFIX_RE = re.compile(
    r"""(?ix)
    [,\s]+
    (?:inc\.?|incorporated|ltd\.?|limited|llc|l\.l\.c\.|llp|l\.l\.p\.|lp|l\.p\.|
       corp\.?|corporation|co\.?|company|plc|gmbh|s\.a\.|s\.r\.l\.|sarl|ag)
    \.?$
    """
)


def _client_short_name(client_legal_name: str) -> str:
    s = (client_legal_name or "").strip()
    # Iteratively strip common legal suffixes.
    prev = None
    while s and s != prev:
        prev = s
        s = _LEGAL_SUFFIX_RE.sub("", s).strip().rstrip(",")
    return s or (client_legal_name or "").strip()


def _normalize_money_value(value) -> str:
    """
    Normalize a money-like input to a plain integer string with commas, without
    any currency symbol.

    Back-compat: if the user provides a small integer (< 1000), assume legacy
    "in thousands" and multiply by 1,000.
    """
    if value is None:
        return ""
    s = str(value).strip()
    m = _MONEY_INT_RE.match(s)
    if not m:
        return s
    digits = m.group(1).replace(",", "")
    try:
        n = int(digits)
    except ValueError:
        return s
    if n < 1000:
        n *= 1000
    return f"{n:,}"


def _normalize_lookup_key(value: str) -> str:
    s = (value or "").strip().lower()
    return re.sub(r"[^a-z0-9]+", "_", s).strip("_")


def resolve_industry_key(industry: str, scope_library: dict) -> tuple[str, Optional[str]]:
    """
    Resolve a user-provided industry key to a supported scope library key.

    Behaviors:
    - exact match
    - normalized/alias match
    - fuzzy typo correction
    - fallback to generic (if present), with warning text
    """
    modules = scope_library.get("industry_modules") if isinstance(scope_library, dict) else None
    if not isinstance(modules, dict) or not modules:
        return industry, None

    supported = sorted(k for k in modules.keys() if isinstance(k, str))
    requested = (industry or "").strip()
    if requested in modules:
        return requested, None

    norm_to_supported = {_normalize_lookup_key(k): k for k in supported}
    alias_map = {_normalize_lookup_key(k): v for k, v in _INDUSTRY_ALIAS_MAP.items()}
    requested_norm = _normalize_lookup_key(requested)

    # "generic" is a valid intentional mode that uses common-skeleton scope only.
    if requested_norm in {"generic", "common"}:
        return "generic", None

    if requested_norm in norm_to_supported:
        resolved = norm_to_supported[requested_norm]
        return resolved, f"Industry '{industry}' normalized to '{resolved}'."

    aliased = alias_map.get(requested_norm)
    if aliased in modules:
        return aliased, f"Industry '{industry}' mapped to '{aliased}'."

    fuzzy_candidates = sorted(set(norm_to_supported.keys()) | set(alias_map.keys()))
    match = difflib.get_close_matches(requested_norm, fuzzy_candidates, n=1, cutoff=0.82)
    if match:
        m = match[0]
        resolved = norm_to_supported.get(m) or alias_map.get(m)
        if resolved in modules:
            return resolved, f"Industry '{industry}' corrected to '{resolved}'."

    if "generic" in modules:
        return "generic", f"Industry '{industry}' not recognized; using 'generic' scope."

    # Even when `generic` is not an explicit industry module key, this still
    # yields common-skeleton-only scope insertion (the intended generic fallback).
    return "generic", f"Industry '{industry}' not recognized; using common-scope fallback."


def _infer_template_type(template_file: str) -> Optional[str]:
    name = Path(template_file).name.lower()
    if "buyside" in name:
        return "buyside"
    if "sellside" in name:
        return "sellside"
    return None


def _iter_table_paragraphs(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                yield para
            for nested_table in cell.tables:
                yield from _iter_table_paragraphs(nested_table)


def _iter_document_paragraphs(doc: Document):
    # Body paragraphs
    for para in doc.paragraphs:
        yield para

    # Body tables (including nested tables)
    for table in doc.tables:
        yield from _iter_table_paragraphs(table)

    # Headers/footers (paragraphs + tables, including nested tables)
    for section in doc.sections:
        for para in section.header.paragraphs:
            yield para
        for table in section.header.tables:
            yield from _iter_table_paragraphs(table)

        for para in section.footer.paragraphs:
            yield para
        for table in section.footer.tables:
            yield from _iter_table_paragraphs(table)


def _extract_placeholders(doc: Document) -> set[str]:
    """
    Extract placeholder keys (without braces) from the current document state.
    Uses python-docx `.text` so split-across-runs placeholders are recognized.
    """
    keys: set[str] = set()

    def _scan_text(text: str) -> None:
        for token in PLACEHOLDER_RE.findall(text or ""):
            inner = token[2:-2].strip()
            if token in GUIDANCE_KEYS or f"{{{{{inner}}}}}" in GUIDANCE_KEYS:
                continue
            if inner:
                keys.add(inner)

    for para in _iter_document_paragraphs(doc):
        _scan_text(para.text)

    return keys


def _default_review(key: str) -> str:
    return f"[REVIEW: Provide {key}]"


def _has_usable_value(out: dict, key: str) -> bool:
    """
    Return True if the dict has an acceptable value for a placeholder key.
    For some keys, empty string is acceptable to intentionally omit a clause.
    """
    if key not in out:
        return False
    if key in ALLOW_EMPTY_KEYS:
        return out.get(key) is not None
    return not _is_blank(out.get(key))


def derive_and_default_variables(
    *,
    template_file: str,
    variables: dict,
    placeholder_keys: set[str],
) -> dict:
    """
    Apply deterministic inference + defaults before gating/replacement.
    Never overrides a non-blank user-provided value.
    """
    out = dict(variables or {})

    # Keep release-letter placeholders as raw template placeholders, even if
    # values are passed in upstream variables.
    for key in PRESERVE_TEMPLATE_PLACEHOLDERS:
        out.pop(key, None)
    template_type = _infer_template_type(template_file)

    # Dates
    _set_if_blank(out, "LETTER_DATE", _format_today())
    _set_if_blank(out, "DATE_COMMENCE", "Immediately")
    _set_if_blank(out, "DATE_DRAFT_DELIVERY", "TBD")
    _set_if_blank(out, "FAL_DATE", "TBD")
    if template_type == "sellside":
        _set_if_blank(out, "FAL_LETTER_DATE", out.get("FAL_DATE"))

    # Signing / names
    _set_if_blank(out, "CLIENT_LEGAL_NAME_FULL", out.get("CLIENT_LEGAL_NAME"))
    _set_if_blank(out, "CLIENT_LEGAL_NAME_ACCEPTANCE", out.get("CLIENT_LEGAL_NAME"))
    if _is_blank(out.get("CLIENT_NAME_SHORT")) and not _is_blank(out.get("CLIENT_LEGAL_NAME")):
        out["CLIENT_NAME_SHORT"] = _client_short_name(str(out.get("CLIENT_LEGAL_NAME")))

    _set_if_blank(out, "CLIENT_CONTACT_NAME_TITLE", out.get("CLIENT_CONTACT_NAME"))
    _set_if_blank(out, "EL_SIGNATORY_NAME", out.get("CLIENT_CONTACT_NAME"))
    _set_if_blank(out, "BILLING_ENTITY_NAME", out.get("CLIENT_LEGAL_NAME"))
    _set_if_blank(out, "CLIENT_COUNTRY_2", out.get("CLIENT_COUNTRY"))

    # Signing partners (often same as lead partner)
    _set_if_blank(out, "SIGNING_PARTNER_NAME", out.get("TEAM_PARTNER_NAME"))

    # Project naming
    if _is_blank(out.get("PROJECT_CODE_NAME")):
        project_name = out.get("PROJECT_NAME")
        if not _is_blank(project_name):
            out["PROJECT_CODE_NAME"] = str(project_name).strip()
        elif not _is_blank(out.get("PROJECT_NAME_HEADER")):
            header = str(out.get("PROJECT_NAME_HEADER")).strip()
            out["PROJECT_CODE_NAME"] = re.sub(r"(?i)^project\s+", "", header).strip() or header

    # Project header
    if _is_blank(out.get("PROJECT_NAME_HEADER")) and not _is_blank(out.get("PROJECT_CODE_NAME")):
        code = str(out.get("PROJECT_CODE_NAME")).strip()
        out["PROJECT_NAME_HEADER"] = code if code.lower().startswith("project ") else f"Project {code}"

    # Deal type variants (buyside)
    # Target descriptions (used in various header/appendix contexts)
    target_legal = out.get("TARGET_LEGAL_NAME_FULL") or out.get("TARGET_LEGAL_NAME")
    _set_if_blank(out, "TARGET_DESCRIPTION", target_legal)
    _set_if_blank(out, "TARGET_DESCRIPTION_DETAIL", target_legal)
    _set_if_blank(out, "CLIENT_AND_TARGET", "Client and Target")

    if template_type == "sellside":
        _set_if_blank(out, "DEAL_DESCRIPTION", "potential divestment of Target")
        _set_if_blank(out, "TARGET_DEFINED_TERM", '("Target")')
        _set_if_blank(out, "CHOICE_CLIENT_OR_TARGET", "Target")
        _set_if_blank(out, "CHOICE_CLIENT_OR_TARGET_POSSESSIVE", "Target's")
        _set_if_blank(out, "CHOICE_CLIENT_OR_TARGET_POSSESSIVE_2", "Target's")
        _set_if_blank(out, "CLIENT_OR_TARGET", "Target")
        _set_if_blank(out, "CLIENT_AND_OR_TARGET", "Target")
        _set_if_blank(out, "CLIENT_AND_BIDDER", "Client and Target")
        _set_if_blank(out, "BIDDER_NAME", "Target")

    # Deal type normalization for buyside placeholders.
    if template_type == "buyside" and _is_blank(out.get("CHOICE_DEAL_TYPE")):
        candidate = str(
            out.get("DEAL_TYPE")
            or out.get("DEAL_DESCRIPTION")
            or ""
        ).strip().lower()
        if "invest" in candidate:
            out["CHOICE_DEAL_TYPE"] = "investment in"
        elif candidate:
            out["CHOICE_DEAL_TYPE"] = "acquisition of"

    # Report format default
    _set_if_blank(out, "CHOICE_REPORT_FORMAT", "a PDF written report and an Excel data book")

    # Ethical divider / confidentiality defaults (conflict-language placeholders).
    _set_if_blank(out, "EXISTING_SERVICES_DESCRIPTION", "KPMG provides Advisory and Tax services")
    _set_if_blank(
        out,
        "EXISTING_SERVICES_CAVEAT",
        "and its affiliates. In order to safeguard against the flow of confidential information between "
        "separate KPMG engagement teams servicing Client and Target, we have established formal Ethical "
        "Dividers and Confidentiality Procedures.",
    )
    _set_if_blank(out, "CONFIDENTIALITY_PROCEDURES", _DEFAULT_CONFIDENTIALITY_PROCEDURES)
    _set_if_blank(out, "INDEMNITY_EXCLUSION", "[REVIEW: Confirm indemnity exclusion (or leave blank)]")
    _set_if_blank(out, "CLOSING_SENTIMENT", "[REVIEW: Insert closing sentiment]")

    # Fee formatting (full dollars). Also supports legacy "in thousands" inputs.
    for key in MONEY_KEYS:
        if not _is_blank(out.get(key)):
            out[key] = _normalize_money_value(out.get(key))

    # Ensure every placeholder present in the template has a value (never leave raw {{...}}).
    for key in placeholder_keys:
        if key in PRESERVE_TEMPLATE_PLACEHOLDERS:
            continue
        if not _has_usable_value(out, key):
            out[key] = "" if key in ALLOW_EMPTY_KEYS else _default_review(key)

    return out


# ---------------------------------------------------------------------------
# Optional section removal (marker-based, with heading fallback)
# ---------------------------------------------------------------------------

def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element  # type: ignore[attr-defined]
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _paragraph_has_section_break(paragraph: Paragraph) -> bool:
    """
    Return True when a paragraph contains section properties (`w:sectPr`).

    These markers define section boundaries (for example, Appendix A -> Appendix B)
    and must be preserved when replacing scope text.
    """
    element = getattr(paragraph, "_element", None)
    if element is None:
        return False
    return bool(element.xpath(".//*[local-name()='sectPr']"))


def _normalize_yes_no(value) -> str:
    s = ("" if value is None else str(value)).strip().lower()
    if s in {"y", "yes", "true", "1"}:
        return "yes"
    if s in {"n", "no", "false", "0"}:
        return "no"
    return ""


def _find_paragraph_index(paragraphs: list[Paragraph], needle: str, start: int = 0) -> int:
    for i in range(start, len(paragraphs)):
        if needle in (paragraphs[i].text or ""):
            return i
    return -1


def apply_independence_section_policy(doc: Document, variables: dict, summary_steps: list[str]) -> None:
    """
    Remove Independence Considerations when it does not apply.

    Preferred: marker-based deletion between BLOCK_INDEPENDENCE_START/END paragraphs.
    Fallback: delete from 'Independence Considerations' heading up to (but not including)
    'Management’s Responsibilities'.
    """
    applies = _normalize_yes_no(variables.get("CHOICE_INDEPENDENCE_APPLIES")) or "yes"

    # Preferred: markers (standalone paragraphs)
    paragraphs = list(doc.paragraphs)
    start_idx = _find_paragraph_index(paragraphs, BLOCK_INDEPENDENCE_START)
    end_idx = _find_paragraph_index(paragraphs, BLOCK_INDEPENDENCE_END, start=max(0, start_idx))

    if start_idx != -1 and end_idx != -1 and end_idx >= start_idx:
        if applies == "no":
            for i in range(end_idx, start_idx - 1, -1):
                _delete_paragraph(paragraphs[i])
            summary_steps.append("Independence Considerations: removed (CHOICE_INDEPENDENCE_APPLIES=no)")
            return

        # applies == "yes": remove marker paragraphs only
        _delete_paragraph(paragraphs[end_idx])
        _delete_paragraph(paragraphs[start_idx])
        summary_steps.append("Independence Considerations: kept (removed block markers)")
        return

    # Fallback: heading-based deletion (only if asked to remove)
    if applies != "no":
        return

    indep_heading = "Independence Considerations"
    mgmt_heading = "Management’s Responsibilities"
    ih = _find_paragraph_index(paragraphs, indep_heading)
    if ih == -1:
        return
    mh = _find_paragraph_index(paragraphs, mgmt_heading, start=ih + 1)
    if mh == -1:
        return

    for i in range(mh - 1, ih - 1, -1):
        _delete_paragraph(paragraphs[i])
    summary_steps.append("Independence Considerations: removed (heading fallback)")


# ---------------------------------------------------------------------------
# Placeholder replacement (run-safe, handles split-across-runs)
# ---------------------------------------------------------------------------

def _replace_in_runs(runs, old: str, new: str) -> int:
    """Replace a placeholder that may span multiple runs. Returns replacement count."""
    count = 0
    # Fast path: single-run match
    for run in runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    if count > 0:
        return count

    # Slow path: placeholder split across runs
    while True:
        full_text = "".join(run.text for run in runs)
        idx = full_text.find(old)
        if idx == -1:
            break
        end_idx = idx + len(old)
        cursor = 0
        inserted = False

        for run in runs:
            text = run.text
            run_end = cursor + len(text)

            if run_end <= idx or cursor >= end_idx:
                cursor = run_end
                continue

            before = text[: max(0, idx - cursor)] if idx > cursor else ""
            after = text[end_idx - cursor:] if run_end > end_idx else ""

            if not inserted:
                run.text = before + str(new) + after
                inserted = True
            else:
                run.text = after

            cursor = run_end

        count += 1

    return count


def _replace_placeholder_key_in_runs(runs, key: str, new: str) -> int:
    """
    Replace placeholder tokens for a key, allowing optional whitespace inside braces.

    Examples matched:
      {{KEY}}, {{ KEY }}, {{ KEY}}, {{KEY }}
    """
    key = (key or "").strip()
    if not key:
        return 0

    pattern = re.compile(r"\{\{\s*" + re.escape(key) + r"\s*\}\}")
    count = 0

    while True:
        full_text = "".join(run.text for run in runs)
        m = pattern.search(full_text)
        if not m:
            break

        idx, end_idx = m.span()
        cursor = 0
        inserted = False

        for run in runs:
            text = run.text
            run_end = cursor + len(text)

            if run_end <= idx or cursor >= end_idx:
                cursor = run_end
                continue

            before = text[: max(0, idx - cursor)] if idx > cursor else ""
            after = text[end_idx - cursor:] if run_end > end_idx else ""

            if not inserted:
                run.text = before + str(new) + after
                inserted = True
            else:
                run.text = after

            cursor = run_end

        count += 1

    return count


def replace_placeholders(doc: Document, variables: dict) -> None:
    """Replace all {{KEY}} placeholders throughout the document."""
    # Build replacement map by placeholder key (without braces), allowing
    # whitespace-padded token variants in templates.
    repl: dict[str, str] = {}
    for key, value in variables.items():
        raw = str(key) if key is not None else ""
        if raw.startswith("{{") and raw.endswith("}}"):
            normalized_key = raw[2:-2].strip()
        else:
            normalized_key = raw.strip()
        if not normalized_key:
            continue
        repl[normalized_key] = str(value) if value is not None else ""

    def _process_paragraph(para):
        if not para.runs:
            return
        for placeholder_key, new in repl.items():
            _replace_placeholder_key_in_runs(para.runs, placeholder_key, new)

    for para in _iter_document_paragraphs(doc):
        _process_paragraph(para)


# ---------------------------------------------------------------------------
# Guidance block removal
# ---------------------------------------------------------------------------

def remove_guidance_blocks(doc: Document) -> int:
    """
    Remove all internal guidance paragraphs from the document.

    Supports both:
    - Numeric placeholders: {{GUIDANCE_NN}}
    - Freeform guidance blocks: {{GUIDANCE: ...}}
    Returns count of paragraphs removed.
    """
    removed = 0

    def _is_guidance(text: str) -> bool:
        return bool(GUIDANCE_RE.search(text) or GUIDANCE_FREEFORM_RE.search(text))

    def _collect_paragraphs_from_table(table, out: list[Paragraph]) -> None:
        for row in table.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
                for nested_table in cell.tables:
                    _collect_paragraphs_from_table(nested_table, out)

    paragraphs: list[Paragraph] = []

    # Body paragraphs
    paragraphs.extend(doc.paragraphs)

    # Tables (body + nested)
    for table in doc.tables:
        _collect_paragraphs_from_table(table, paragraphs)

    # Headers/footers (paragraphs + tables)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        for table in section.header.tables:
            _collect_paragraphs_from_table(table, paragraphs)

        paragraphs.extend(section.footer.paragraphs)
        for table in section.footer.tables:
            _collect_paragraphs_from_table(table, paragraphs)

    to_remove: list[Paragraph] = []
    seen = set()
    for para in paragraphs:
        element = getattr(para, "_element", None)
        if element is None:
            continue
        element_id = id(element)
        if element_id in seen:
            continue
        seen.add(element_id)
        if _is_guidance(para.text or ""):
            to_remove.append(para)

    for para in to_remove:
        _delete_paragraph(para)
        removed += 1

    return removed


# ---------------------------------------------------------------------------
# FDD Scope block replacement
# ---------------------------------------------------------------------------

_DIRECTIVE_VERBS = {
    "obtain",
    "read",
    "review",
    "understand",
    "analyze",
    "perform",
    "consider",
    "identify",
    "determine",
    "develop",
    "discuss",
    "evaluate",
    "compare",
    "summarize",
    "comment",
    "meet",
    "inquire",
    "assess",
    "bridge",
    "reconcile",
    "gain",
    "propose",
    "segment",
    "calculate",
    "collect",
    "confirm",
    "test",
    "prepare",
    "provide",
}


def _first_word(text: str) -> str:
    m = re.match(r"[\W_]*(\w+)", text.strip())
    return (m.group(1) if m else "").lower()


def _looks_like_new_directive(text: str) -> bool:
    w = _first_word(text)
    return w in _DIRECTIVE_VERBS and len(text.strip().split()) >= 3


@dataclass
class _ScopeNode:
    text: str
    children: list["_ScopeNode"] = field(default_factory=list)
    top_level_id: Optional[str] = None


def _is_v2_bullet_list(bullets: list) -> bool:
    """Return True if bullets appear to be v2 objects with text/children."""
    if not isinstance(bullets, list) or not bullets:
        return False
    sample = bullets[0]
    return isinstance(sample, dict) and isinstance(sample.get("text"), str)


def _parse_scope_bullets(bullets: list) -> list[_ScopeNode]:
    """
    Convert a flat list of scope bullet strings into a simple parent/child tree.

    Heuristic:
    - Lines ending with ':' become parents; following lines attach beneath them.
    - A “new directive” line (e.g., starts with obtain/analyze/perform/consider/...) ends the current nesting.
    """
    root: list[_ScopeNode] = []
    stack: list[_ScopeNode] = []

    def _add(text: str) -> _ScopeNode:
        node = _ScopeNode(text=text)
        if stack:
            stack[-1].children.append(node)
        else:
            root.append(node)
        return node

    for bullet in bullets:
        if not isinstance(bullet, str):
            continue
        text = bullet.strip()
        if not text:
            continue

        if text.endswith(":"):
            if _looks_like_new_directive(text):
                stack.clear()
            node = _add(text)
            stack.append(node)
            continue

        if stack and _looks_like_new_directive(text):
            stack.clear()
        _add(text)

    return root


def _scope_nodes_from_v2(*, bullets: list) -> list[_ScopeNode]:
    """
    Build scope nodes from v2 bullet objects.

    v2 node shape:
      {"id": "scope.001", "text": "...", "children": [ ... ]}
    """
    def build(node: dict) -> _ScopeNode:
        text = node.get("text", "") if isinstance(node, dict) else ""
        n = _ScopeNode(text=text, children=[], top_level_id=node.get("id") if isinstance(node, dict) else None)
        for child in (node.get("children") or []) if isinstance(node, dict) else []:
            if isinstance(child, dict):
                n.children.append(build(child))
        return n

    out: list[_ScopeNode] = []
    for node in bullets:
        if isinstance(node, dict) and isinstance(node.get("text"), str):
            out.append(build(node))
    return out


def _scope_nodes_from_schema(*, bullets: list, schema_nodes: list) -> list[_ScopeNode]:
    """
    Build scope nodes from explicit schema (indices into the bullets list).

    Schema node shape:
      {"i": <index>, "id": "scope.xxxxxxxx", "children": [ ... ]}
    """
    def build(node: dict) -> _ScopeNode:
        i = int(node["i"])
        text = bullets[i] if 0 <= i < len(bullets) and isinstance(bullets[i], str) else ""
        n = _ScopeNode(text=text, children=[], top_level_id=node.get("id"))
        for child in (node.get("children") or []):
            if isinstance(child, dict):
                n.children.append(build(child))
        return n

    out: list[_ScopeNode] = []
    for node in schema_nodes:
        if isinstance(node, dict):
            out.append(build(node))
    return out


def _assign_top_level_ids(nodes: list[_ScopeNode], start: int) -> int:
    """
    Assign stable ids to top-level scope nodes in document insertion order.

    IDs are global (not section-specific) to avoid exposing internal section keys in Canvas.
    """
    n = start
    for node in nodes:
        if not node.top_level_id:
            node.top_level_id = f"scope.{n:03d}"
        n += 1
    return n


def _indent_for_depth(depth: int) -> Emu:
    if depth <= 0:
        return _INDENT_BULLET
    if depth == 1:
        return _INDENT_SUBBULLET
    return _INDENT_SUBSUBBULLET


def _get_paragraph_num_info(paragraph: Paragraph) -> tuple[Optional[int], Optional[int]]:
    ppr = paragraph._p.pPr
    if ppr is None or ppr.numPr is None:
        return None, None
    num_id = ppr.numPr.numId.val if ppr.numPr.numId is not None else None
    ilvl = ppr.numPr.ilvl.val if ppr.numPr.ilvl is not None else 0
    try:
        return (int(num_id) if num_id is not None else None, int(ilvl) if ilvl is not None else 0)
    except (TypeError, ValueError):
        return None, None


def _extract_numbering_metadata(doc: Document) -> tuple[dict[int, int], dict[int, tuple[Optional[str], Optional[str]]]]:
    numbering = doc.part.numbering_part.numbering_definitions._numbering

    num_to_abstract: dict[int, int] = {}
    for num in numbering.num_lst:
        try:
            num_to_abstract[int(num.numId)] = int(num.abstractNumId.val)
        except (TypeError, ValueError):
            continue

    abstract_level0: dict[int, tuple[Optional[str], Optional[str]]] = {}
    for abstract in numbering.findall(qn("w:abstractNum")):
        abstract_id_raw = abstract.get(qn("w:abstractNumId"))
        try:
            abstract_id = int(abstract_id_raw) if abstract_id_raw is not None else None
        except ValueError:
            abstract_id = None
        if abstract_id is None:
            continue

        level0 = None
        for level in abstract.findall(qn("w:lvl")):
            if level.get(qn("w:ilvl")) == "0":
                level0 = level
                break
        if level0 is None:
            continue

        fmt_el = level0.find(qn("w:numFmt"))
        text_el = level0.find(qn("w:lvlText"))
        abstract_level0[abstract_id] = (
            fmt_el.get(qn("w:val")) if fmt_el is not None else None,
            text_el.get(qn("w:val")) if text_el is not None else None,
        )

    return num_to_abstract, abstract_level0


def _find_abstract_id(
    abstract_level0: dict[int, tuple[Optional[str], Optional[str]]],
    *,
    num_fmt: str,
    lvl_text_suffix: Optional[str] = None,
) -> Optional[int]:
    for abstract_id in sorted(abstract_level0.keys()):
        fmt, text = abstract_level0[abstract_id]
        if fmt != num_fmt:
            continue
        if lvl_text_suffix and not str(text or "").endswith(lvl_text_suffix):
            continue
        return abstract_id
    return None


def _set_paragraph_numbering(paragraph: Paragraph, *, num_id: Optional[int], ilvl: int = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    if ppr.numPr is not None:
        ppr._remove_numPr()

    if num_id is None:
        return

    num_pr = ppr.get_or_add_numPr()
    num_pr.get_or_add_numId().val = int(num_id)
    num_pr.get_or_add_ilvl().val = int(ilvl)


def load_scope_library(path: str) -> dict:
    """Load the FDD scope library JSON."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_scope_review_buckets(path: str) -> dict:
    """Load scope-review-buckets JSON used for grouping/ordering (best-effort)."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_section_applicability(path: str) -> dict:
    """Load section-applicability JSON used for runtime filtering/replacements."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_optional_scope_library(path: str) -> dict:
    """Load optional scope library JSON used for explicit opt-in scope additions."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _load_scope_buckets_for_scope_library(scope_library_file: str) -> Optional[dict]:
    """
    Best-effort companion loader.

    If `scope-review-buckets.json` is present beside the scope library file, use it.
    Generation must remain robust when this file is absent.
    """
    try:
        lib_path = Path(scope_library_file).resolve()
        candidate = lib_path.with_name("scope-review-buckets.json")
        if candidate.exists():
            return load_scope_review_buckets(str(candidate))
    except Exception:
        return None
    return None


def _load_section_applicability_for_scope_library(scope_library_file: str) -> Optional[dict]:
    """
    Best-effort applicability loader.

    Priority:
    1) Dist companion file beside scope-library.json
    2) Repo docs path (for local development workflows)
    """
    try:
        lib_path = Path(scope_library_file).resolve()
        candidates = [
            lib_path.with_name("section-applicability.json"),
            lib_path.parent.parent / "docs" / "scope-library" / "section-applicability.json",
        ]
        for candidate in candidates:
            if candidate.exists():
                return load_section_applicability(str(candidate))
    except Exception:
        return None
    return None


def _load_optional_scope_library_for_scope_library(scope_library_file: str) -> Optional[dict]:
    """
    Best-effort optional scope loader.

    Priority:
    1) Dist companion file beside scope-library.json
    2) Project dist file
    """
    try:
        lib_path = Path(scope_library_file).resolve()
        candidates = [
            lib_path.with_name("scope-library-optional.json"),
            lib_path.parent.parent / "dist" / "scope-library-optional.json",
        ]
        for candidate in candidates:
            if candidate.exists():
                return load_optional_scope_library(str(candidate))
    except Exception:
        return None
    return None


def _apply_optional_scope_modules(
    *,
    common_skeleton: list[dict],
    industry_module: dict[str, Any],
    industry: str,
    optional_library: Optional[dict],
    optional_section_keys: set[str],
    ad_hoc_optional_sections: dict[str, list[dict]],
) -> dict:
    return _core_apply_optional_scope_modules(
        common_skeleton=common_skeleton,
        industry_module=industry_module,
        industry=industry,
        optional_library=optional_library,
        optional_section_keys=optional_section_keys,
        ad_hoc_optional_sections=ad_hoc_optional_sections,
        enable_unknown_fallback=True,
    )


def _build_industry_scope_view(
    *,
    scope_library: dict,
    industry: str,
    applicability: Optional[dict],
) -> dict:
    return _core_build_industry_scope_view(
        scope_library=scope_library,
        industry=industry,
        applicability=applicability,
        include_summary=False,
    )


def _ordered_active_section_keys(
    *,
    common_keys_in_order: list[str],
    module_keys_in_order: list[str],
    excluded_section_keys: set[str],
    scope_buckets: Optional[dict],
) -> list[str]:
    return _core_ordered_active_section_keys(
        common_keys_in_order=common_keys_in_order,
        module_keys_in_order=module_keys_in_order,
        excluded_section_keys=excluded_section_keys,
        scope_buckets=scope_buckets,
    )


def replace_fdd_scope_block(
    doc: Document,
    industry: str,
    scope_library: dict,
    scope_selection: Optional[dict] = None,
    scope_heading_search: str = "FINANCIAL DUE DILIGENCE",
    end_boundary_search: str = "These Terms and Conditions",
) -> dict:
    """
    Replace the sample FDD scope block with industry-specific content.

    Returns summary dict with counts.
    """
    paras = list(doc.paragraphs)

    # Find scope boundaries
    start_idx = None
    end_idx = None
    for i, p in enumerate(paras):
        if start_idx is None and scope_heading_search in p.text:
            start_idx = i
        elif start_idx is not None and p.text.strip().startswith(end_boundary_search):
            end_idx = i
            break

    if start_idx is None:
        return {"error": f"Scope heading '{scope_heading_search}' not found", "sections_inserted": 0, "bullets_inserted": 0}
    if end_idx is None:
        return {"error": f"End boundary '{end_boundary_search}' not found", "sections_inserted": 0, "bullets_inserted": 0}

    # Capture formatting templates + numbering archetypes from the sample scope.
    # This keeps generated scope numbering stable and dynamic even when items are
    # added/removed (e.g., exclusions).
    num_to_abstract, abstract_level0 = _extract_numbering_metadata(doc)

    heading_template = None
    alpha_template = None
    roman_template = None
    fallback_nonbold_template = None

    section_abstract_id: Optional[int] = None
    alpha_abstract_id: Optional[int] = None
    roman_abstract_id: Optional[int] = None

    sample_paras = [p for p in paras[start_idx + 1: end_idx] if p.text.strip()]
    for p in sample_paras:
        num_id, _ = _get_paragraph_num_info(p)
        abstract_id = num_to_abstract.get(num_id) if num_id is not None else None
        fmt, lvl_text = abstract_level0.get(abstract_id, (None, None)) if abstract_id is not None else (None, None)

        if fallback_nonbold_template is None and not any(r.bold for r in p.runs):
            fallback_nonbold_template = p

        if heading_template is None and any(r.bold for r in p.runs) and fmt == "decimal":
            heading_template = p
            section_abstract_id = abstract_id

        if alpha_template is None and fmt == "lowerLetter" and str(lvl_text or "").endswith(")"):
            alpha_template = p
            alpha_abstract_id = abstract_id

        if roman_template is None and fmt == "lowerRoman":
            roman_template = p
            roman_abstract_id = abstract_id

    if heading_template is None:
        for p in sample_paras:
            if any(r.bold for r in p.runs):
                heading_template = p
                num_id, _ = _get_paragraph_num_info(p)
                section_abstract_id = num_to_abstract.get(num_id) if num_id is not None else section_abstract_id
                break
    if heading_template is None and sample_paras:
        heading_template = sample_paras[0]

    if alpha_template is None:
        for p in sample_paras:
            num_id, _ = _get_paragraph_num_info(p)
            abstract_id = num_to_abstract.get(num_id) if num_id is not None else None
            fmt, _ = abstract_level0.get(abstract_id, (None, None)) if abstract_id is not None else (None, None)
            if fmt == "lowerLetter":
                alpha_template = p
                alpha_abstract_id = abstract_id
                break
    if alpha_template is None:
        alpha_template = fallback_nonbold_template or heading_template

    if roman_template is None:
        roman_template = alpha_template

    # Fallback abstracts when a direct template match isn't available.
    if section_abstract_id is None:
        section_abstract_id = _find_abstract_id(abstract_level0, num_fmt="decimal", lvl_text_suffix=".")
    if alpha_abstract_id is None:
        alpha_abstract_id = _find_abstract_id(abstract_level0, num_fmt="lowerLetter", lvl_text_suffix=")")
    if alpha_abstract_id is None:
        alpha_abstract_id = _find_abstract_id(abstract_level0, num_fmt="lowerLetter")
    if roman_abstract_id is None:
        roman_abstract_id = _find_abstract_id(abstract_level0, num_fmt="lowerRoman", lvl_text_suffix=".")
    if roman_abstract_id is None:
        roman_abstract_id = _find_abstract_id(abstract_level0, num_fmt="lowerRoman")

    if heading_template is None:
        return {"error": "No formatting template found in scope block", "sections_inserted": 0, "bullets_inserted": 0}

    # Delete scope content, but preserve section-break paragraphs so Appendix
    # boundaries/headers remain intact (Appendix A scope vs Appendix B terms).
    to_remove = paras[start_idx + 1: end_idx]
    preserved_section_breaks = 0
    for p in to_remove:
        if _paragraph_has_section_break(p):
            preserved_section_breaks += 1
            continue
        p._p.getparent().remove(p._p)

    # Build and insert replacement content (runtime applicability-aware)
    scope_view = _build_industry_scope_view(
        scope_library=scope_library,
        industry=industry,
        applicability=scope_library.get("_section_applicability"),
    )
    common = scope_view.get("common_skeleton", []) or []
    modules = scope_view.get("industry_module", {}) or {}

    parsed_scope = _core_parse_scope_selection(scope_selection if isinstance(scope_selection, dict) else None)
    excluded_top_level_ids = set(parsed_scope.get("excluded_top_level_ids") or [])
    excluded_section_keys = set(parsed_scope.get("excluded_section_keys") or [])
    optional_section_keys = set(parsed_scope.get("optional_section_keys") or [])
    ad_hoc_optional_sections = parsed_scope.get("ad_hoc_optional_sections") or {}
    ad_hoc_sections_normalized_from_list = parsed_scope.get("ad_hoc_sections_normalized_from_list") or []
    ad_hoc_sections_invalid_items = parsed_scope.get("ad_hoc_sections_invalid_items") or []

    optional_summary = _apply_optional_scope_modules(
        common_skeleton=common,
        industry_module=modules,
        industry=industry,
        optional_library=scope_library.get("_optional_scope_library"),
        optional_section_keys=optional_section_keys,
        ad_hoc_optional_sections=ad_hoc_optional_sections,
    )

    common_by_key = {
        s.get("normalized_heading"): s
        for s in common
        if isinstance(s, dict) and isinstance(s.get("normalized_heading"), str)
    }

    cursor = paras[start_idx]
    section_count = 0
    bullet_count = 0
    next_scope_id = 1

    scope_schema = scope_library.get("scope_schema") if isinstance(scope_library, dict) else None
    schema_nesting = scope_schema.get("nesting") if isinstance(scope_schema, dict) else None
    common_schema = schema_nesting.get("common_skeleton") if isinstance(schema_nesting, dict) else None
    industry_schema = schema_nesting.get("industry_modules") if isinstance(schema_nesting, dict) else None

    numbering = doc.part.numbering_part.numbering_definitions._numbering
    heading_num_fallback_id, _ = _get_paragraph_num_info(heading_template)
    alpha_num_fallback_id, _ = _get_paragraph_num_info(alpha_template)
    roman_num_fallback_id, _ = _get_paragraph_num_info(roman_template)

    def _new_num_id(
        abstract_id: Optional[int],
        fallback_num_id: Optional[int] = None,
        *,
        restart_at: Optional[int] = 1,
    ) -> Optional[int]:
        if abstract_id is not None:
            try:
                num = numbering.add_num(int(abstract_id))
                if restart_at is not None:
                    lvl_override = num.add_lvlOverride(0)
                    lvl_override.add_startOverride(int(restart_at))
                return int(num.numId)
            except Exception:
                pass
        return fallback_num_id

    section_num_id = _new_num_id(section_abstract_id, heading_num_fallback_id)

    def _insert_after(cursor_para, text, *, template_para: Paragraph, bold=False, indent=None, num_id: Optional[int] = None, ilvl: int = 0):
        new_p = deepcopy(template_para._p)
        cursor_para._p.addnext(new_p)
        new_para = Paragraph(new_p, cursor_para._parent)

        if new_para.runs:
            new_para.runs[0].text = text
            for run in new_para.runs[1:]:
                run.text = ""
            for run in new_para.runs:
                run.bold = bold
        else:
            run = new_para.add_run(text)
            run.bold = bold

        _set_paragraph_numbering(new_para, num_id=num_id, ilvl=ilvl)

        pf = new_para.paragraph_format
        pf.left_indent = indent

        try:
            new_para.style = doc.styles["Body Text"]
        except KeyError:
            pass

        return new_para

    def _insert_nodes(nodes: list[_ScopeNode], *, alpha_num_id: Optional[int]) -> None:
        nonlocal cursor, bullet_count

        def _insert_children(children: list[_ScopeNode], depth: int, *, roman_num_id: Optional[int]) -> None:
            nonlocal cursor, bullet_count
            for child in children:
                # Cap visual depth so output stays at the agreed 3-tier model:
                # section heading (1.) -> parent (a)) -> descendant (i.).
                flattened_depth = 1 if depth >= 1 else depth
                cursor = _insert_after(
                    cursor,
                    child.text,
                    template_para=roman_template,
                    indent=_indent_for_depth(flattened_depth),
                    num_id=roman_num_id,
                )
                bullet_count += 1
                if child.children:
                    _insert_children(child.children, depth + 1, roman_num_id=roman_num_id)

        for node in nodes:
            if node.top_level_id and node.top_level_id in excluded_top_level_ids:
                continue

            cursor = _insert_after(
                cursor,
                node.text,
                template_para=alpha_template,
                indent=_indent_for_depth(0),
                num_id=alpha_num_id,
            )
            bullet_count += 1

            if node.children:
                roman_num_id = _new_num_id(roman_abstract_id, roman_num_fallback_id)
                _insert_children(node.children, 1, roman_num_id=roman_num_id)

    ordered_section_keys = _ordered_active_section_keys(
        common_keys_in_order=[
            s.get("normalized_heading")
            for s in common
            if isinstance(s, dict) and isinstance(s.get("normalized_heading"), str)
        ],
        module_keys_in_order=[
            k for k in modules.keys()
            if isinstance(k, str)
        ],
        excluded_section_keys=excluded_section_keys,
        scope_buckets=scope_library.get("_scope_review_buckets"),
    )

    for section_key in ordered_section_keys:
        normalized_key = _normalize_lookup_key(section_key)
        if normalized_key in excluded_section_keys:
            continue

        common_section = common_by_key.get(section_key)
        if isinstance(common_section, dict):
            heading = common_section.get("heading") or _format_heading_from_key(section_key)
            default_bullets = common_section.get("default_bullets", [])
        else:
            heading = _format_heading_from_key(section_key)
            default_bullets = []
        industry_bullets = modules.get(section_key, [])

        nodes_default = []
        if default_bullets:
            if _is_v2_bullet_list(default_bullets):
                nodes_default = _scope_nodes_from_v2(bullets=default_bullets)
            elif isinstance(common_schema, dict):
                sec = common_schema.get(section_key)
                if isinstance(sec, dict):
                    sn = sec.get("default_bullets")
                    if isinstance(sn, list):
                        nodes_default = _scope_nodes_from_schema(bullets=default_bullets, schema_nodes=sn)
            if not nodes_default:
                nodes_default = _parse_scope_bullets(default_bullets)
            next_scope_id = _assign_top_level_ids(nodes_default, next_scope_id)

        nodes_industry = []
        if industry_bullets:
            if _is_v2_bullet_list(industry_bullets):
                nodes_industry = _scope_nodes_from_v2(bullets=industry_bullets)
            elif isinstance(industry_schema, dict):
                ind = industry_schema.get(industry)
                if isinstance(ind, dict):
                    sn = ind.get(section_key)
                    if isinstance(sn, list):
                        nodes_industry = _scope_nodes_from_schema(
                            bullets=industry_bullets,
                            schema_nodes=sn,
                        )
            if not nodes_industry:
                nodes_industry = _parse_scope_bullets(industry_bullets)
            next_scope_id = _assign_top_level_ids(nodes_industry, next_scope_id)

        nodes = nodes_default + nodes_industry
        if not nodes:
            continue
        if not any(n.top_level_id not in excluded_top_level_ids for n in nodes):
            continue

        alpha_num_id = _new_num_id(alpha_abstract_id, alpha_num_fallback_id)
        cursor = _insert_after(
            cursor,
            heading,
            template_para=heading_template,
            bold=True,
            num_id=section_num_id,
        )
        section_count += 1
        _insert_nodes(nodes, alpha_num_id=alpha_num_id)

    return {
        "sections_inserted": section_count,
        "bullets_inserted": bullet_count,
        "industry": industry,
        "section_breaks_preserved": preserved_section_breaks,
        "excluded_section_keys_applied": sorted(excluded_section_keys),
        "optional_section_keys_requested": optional_summary["optional_section_keys_requested"],
        "optional_sections_applied_common": optional_summary["optional_sections_applied_common"],
        "optional_sections_applied_industry": optional_summary["optional_sections_applied_industry"],
        "optional_sections_unknown": optional_summary["optional_sections_unknown"],
        "optional_sections_synthesized": optional_summary.get("optional_sections_synthesized", []),
        "ad_hoc_sections_applied": optional_summary["ad_hoc_sections_applied"],
        "ad_hoc_sections_normalized_from_list": ad_hoc_sections_normalized_from_list,
        "ad_hoc_sections_invalid_items": ad_hoc_sections_invalid_items,
        "optional_bullets_added": optional_summary["optional_bullets_added"],
    }


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def validate_no_placeholders(doc: Document) -> list:
    """Return list of any remaining {{...}} placeholders."""
    remaining = []

    def _scan(para):
        matches = PLACEHOLDER_RE.findall(para.text)
        for match in matches:
            inner = match[2:-2].strip()
            if inner in PRESERVE_TEMPLATE_PLACEHOLDERS:
                continue
            remaining.append(match)

    for para in _iter_document_paragraphs(doc):
        _scan(para)

    return sorted(set(remaining))


# ---------------------------------------------------------------------------
# Main generation pipeline
# ---------------------------------------------------------------------------

def generate_engagement_letter(
    template_file: str,
    scope_library_file: str,
    industry: str,
    variables: dict,
    scope_selection: Optional[dict] = None,
    output_file: str = "engagement_letter_output.docx",
) -> dict:
    """
    Generate a completed engagement letter.

    Args:
        template_file: Path to the .docx template.
        scope_library_file: Path to the scope library JSON (e.g. scope-library.json).
        industry: Industry key (e.g. "healthcare").
        variables: Dict of KEY -> value (without {{ }} wrappers).
        output_file: Output filename.

    Returns:
        Summary dict with generation results.
    """
    summary = {
        "template": template_file,
        "industry": industry,
        "output": output_file,
        "steps": [],
        "success": False,
    }

    # Step 1: Load template
    doc = Document(template_file)
    summary["steps"].append(f"Loaded template: {template_file}")

    # Step 1.5: Optional section removal (before placeholder extraction)
    apply_independence_section_policy(doc, variables or {}, summary["steps"])

    # Step 2: Remove guidance blocks
    guidance_removed = remove_guidance_blocks(doc)
    summary["steps"].append(f"Removed {guidance_removed} guidance block paragraphs")

    # Step 3: Replace FDD scope block
    scope_library = load_scope_library(scope_library_file)
    scope_buckets = _load_scope_buckets_for_scope_library(scope_library_file)
    if isinstance(scope_buckets, dict):
        scope_library["_scope_review_buckets"] = scope_buckets
        summary["steps"].append("Loaded scope-review-buckets ordering config")
    section_applicability = _load_section_applicability_for_scope_library(scope_library_file)
    if isinstance(section_applicability, dict):
        scope_library["_section_applicability"] = section_applicability
        summary["steps"].append("Loaded section-applicability runtime config")
    optional_scope_library = _load_optional_scope_library_for_scope_library(scope_library_file)
    if isinstance(optional_scope_library, dict):
        scope_library["_optional_scope_library"] = optional_scope_library
        summary["steps"].append("Loaded optional scope runtime config")
    resolved_industry, industry_note = resolve_industry_key(industry, scope_library)
    if industry_note:
        summary["steps"].append(industry_note)
    summary["requested_industry"] = industry
    summary["industry"] = resolved_industry

    scope_result = replace_fdd_scope_block(
        doc,
        industry=resolved_industry,
        scope_library=scope_library,
        scope_selection=scope_selection,
    )
    if "error" in scope_result:
        summary["steps"].append(f"Scope replacement warning: {scope_result['error']}")
    else:
        scope_step = (
            f"Inserted FDD scope: {scope_result['sections_inserted']} sections, "
            f"{scope_result['bullets_inserted']} bullets"
        )
        preserved_breaks = scope_result.get("section_breaks_preserved")
        if isinstance(preserved_breaks, int) and preserved_breaks > 0:
            scope_step += f"; preserved {preserved_breaks} section break marker(s)"
        summary["steps"].append(scope_step)
        optional_common = scope_result.get("optional_sections_applied_common") or []
        optional_industry = scope_result.get("optional_sections_applied_industry") or []
        ad_hoc_applied = scope_result.get("ad_hoc_sections_applied") or []
        optional_added = scope_result.get("optional_bullets_added")
        if optional_common or optional_industry or ad_hoc_applied:
            summary["steps"].append(
                "Applied optional scope sections: "
                f"common={optional_common}, industry={optional_industry}, ad_hoc={ad_hoc_applied}"
            )
        if isinstance(optional_added, int) and optional_added > 0:
            summary["steps"].append(f"Added {optional_added} optional top-level bullet(s)")
        optional_unknown = scope_result.get("optional_sections_unknown") or []
        optional_synth = scope_result.get("optional_sections_synthesized") or []
        ad_hoc_normalized = scope_result.get("ad_hoc_sections_normalized_from_list") or []
        ad_hoc_invalid = scope_result.get("ad_hoc_sections_invalid_items") or []
        if optional_synth:
            summary["steps"].append(
                f"Synthesized ad hoc optional section(s) for unknown key(s): {optional_synth}"
            )
        unresolved_unknown = [k for k in optional_unknown if k not in set(optional_synth)]
        if unresolved_unknown:
            summary["steps"].append(f"Ignored unknown optional section key(s): {unresolved_unknown}")
        if ad_hoc_normalized:
            summary["steps"].append(
                f"Normalized ad hoc optional sections from list payload: {ad_hoc_normalized}"
            )
        if ad_hoc_invalid:
            summary["steps"].append(
                f"Ignored malformed ad hoc optional section entries: {ad_hoc_invalid}"
            )

    # Step 4: Derive defaults + ensure template placeholder coverage (pre-gate)
    placeholder_keys = _extract_placeholders(doc)
    variables = derive_and_default_variables(
        template_file=template_file,
        variables=variables,
        placeholder_keys=placeholder_keys,
    )

    missing_must_ask = [
        k for k in sorted(MUST_ASK_KEYS & placeholder_keys)
        if _is_blank(variables.get(k)) or str(variables.get(k)).startswith(_REVIEW_PREFIX)
    ]
    if missing_must_ask:
        summary["steps"].append(f"ERROR: Missing required user-provided fields: {missing_must_ask}")
        summary["missing_required_fields"] = missing_must_ask
        summary["success"] = False
        return summary

    # Conditional must-ask: SEC/audit status is required only if independence section is present.
    if "CHOICE_SEC_STATUS" in placeholder_keys:
        applies = _normalize_yes_no(variables.get("CHOICE_INDEPENDENCE_APPLIES")) or "yes"
        if applies == "yes" and (
            _is_blank(variables.get("CHOICE_SEC_STATUS"))
            or str(variables.get("CHOICE_SEC_STATUS")).startswith(_REVIEW_PREFIX)
        ):
            summary["steps"].append("ERROR: Missing required user-provided fields: ['CHOICE_SEC_STATUS']")
            summary["missing_required_fields"] = ["CHOICE_SEC_STATUS"]
            summary["success"] = False
            return summary

    summary["steps"].append(f"Derived defaults + filled template placeholders ({len(placeholder_keys)} keys)")

    # Step 5: Replace all field/choice placeholders
    replace_placeholders(doc, variables)
    summary["steps"].append(f"Replaced placeholders ({len(variables)} variables)")

    # Step 6: Validate
    remaining = validate_no_placeholders(doc)
    if remaining:
        summary["steps"].append(f"WARNING: {len(remaining)} unreplaced placeholders: {remaining}")
        summary["remaining_placeholders"] = remaining
    else:
        summary["steps"].append("Validation passed: no remaining placeholders")

    # Step 7: Save
    doc.save(output_file)
    summary["steps"].append(f"Saved: {output_file}")
    summary["success"] = len(remaining) == 0

    return summary


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    import argparse

    def _path_exists_safe(value: str) -> bool:
        """Return True when value is a valid existing path, without raising OSError."""
        try:
            return Path(value).exists()
        except OSError:
            return False

    def _candidate_read_paths(value: str) -> list[Path]:
        """Build likely filesystem candidates for an input path value."""
        raw = str(value or "").strip()
        p = Path(raw)
        out: list[Path] = []

        def _add(path_obj: Path) -> None:
            try:
                normalized = Path(path_obj)
            except Exception:
                return
            if normalized not in out:
                out.append(normalized)

        if p.is_absolute():
            _add(p)
            return out

        script_dir = Path(__file__).resolve().parent
        _add(Path.cwd() / p)
        _add(script_dir / p)
        _add(Path("/mnt/data") / p)
        # Bare filename fallback (helpful when caller passes only basename).
        _add(script_dir / p.name)
        _add(Path("/mnt/data") / p.name)
        return out

    def _resolve_existing_read_path(
        value: str,
        *,
        arg_name: str,
        purpose: str,
        required: bool = True,
    ) -> Optional[str]:
        """Resolve an input path robustly across cwd, script dir, and /mnt/data."""
        raw = str(value or "").strip()
        if not raw:
            if required:
                raise SystemExit(f"Missing required argument: {arg_name}")
            return None

        candidates = _candidate_read_paths(raw)
        for candidate in candidates:
            if _path_exists_safe(str(candidate)):
                try:
                    return str(candidate.resolve())
                except Exception:
                    return str(candidate)

        if required:
            tried = "\n".join(f"  - {c}" for c in candidates) or "  - (no valid candidates)"
            raise SystemExit(
                f"Invalid {arg_name}: {purpose} not found.\n"
                f"  Received: {value}\n"
                f"  CWD: {Path.cwd()}\n"
                f"  Tried:\n{tried}"
            )
        return None

    def _load_json_arg(value: str, *, arg_name: str, required: bool = True):
        """
        Load JSON from either an inline JSON string or a filesystem path.

        Parsing order is intentional:
        1) Try inline JSON first (avoids Path.exists OSError on long JSON strings).
        2) Try path-based JSON.
        """
        if value is None:
            if required:
                raise SystemExit(f"Missing required argument: {arg_name}")
            return None

        # Prefer inline JSON first to avoid path checks on long payloads.
        try:
            return json.loads(value)
        except (json.JSONDecodeError, TypeError):
            pass

        resolved_path = _resolve_existing_read_path(
            value,
            arg_name=arg_name,
            purpose="JSON file",
            required=False,
        )
        if resolved_path:
            with open(resolved_path, "r", encoding="utf-8") as f:
                return json.load(f)

        raise SystemExit(
            f"Invalid {arg_name}: expected a JSON string or path to a JSON file."
        )

    # Back-compat: Code Interpreter sometimes invokes this script with positional
    # args instead of flags. Support both:
    #   python engagement_letter_generator.py --template ... --scope-library ... --industry ... --variables ... --scope-selection ... --output ...
    #   python engagement_letter_generator.py TEMPLATE_FILE SCOPE_LIBRARY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]
    #   python engagement_letter_generator.py TEMPLATE_FILE SCOPE_LIBRARY INDUSTRY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]
    raw_argv = sys.argv[1:]
    has_flag = any(a.startswith("--") for a in raw_argv)

    if has_flag:
        parser = argparse.ArgumentParser(description="Generate an engagement letter")
        parser.add_argument("--template", required=True, help="Path to .docx template")
        parser.add_argument("--scope-library", required=True, help="Path to the scope library JSON (e.g. scope-library.json)")
        parser.add_argument("--industry", required=True, help="Industry key")
        parser.add_argument("--variables", required=True, help="JSON string or path to JSON file with variables")
        parser.add_argument("--scope-selection", help="Optional JSON string or path to JSON file controlling scope inclusion")
        parser.add_argument("--output", default="engagement_letter_output.docx", help="Output filename")
        args = parser.parse_args()
    else:
        if len(raw_argv) < 4:
            raise SystemExit(
                "Usage:\n"
                "  engagement_letter_generator.py --template ... --scope-library ... --industry ... --variables ... [--scope-selection ...] [--output ...]\n"
                "  engagement_letter_generator.py TEMPLATE_FILE SCOPE_LIBRARY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]\n"
                "  engagement_letter_generator.py TEMPLATE_FILE SCOPE_LIBRARY INDUSTRY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]"
            )

        template_file = raw_argv[0]
        scope_library_file = raw_argv[1]
        rest = raw_argv[2:]

        def _is_jsonish(s: str) -> bool:
            s2 = (s or "").lstrip()
            return s2.startswith("{") or s2.startswith("[")

        industry = None
        if len(rest) >= 3 and (not _is_jsonish(rest[0])) and (_path_exists_safe(rest[1]) or _is_jsonish(rest[1])):
            # TEMPLATE SCOPE INDUSTRY VARIABLES SCOPE_SELECTION [OUTPUT]
            industry = rest[0]
            variables_arg = rest[1]
            scope_selection_arg = rest[2]
            output_file = rest[3] if len(rest) >= 4 else "engagement_letter_output.docx"
        else:
            # TEMPLATE SCOPE VARIABLES SCOPE_SELECTION [OUTPUT]
            variables_arg = rest[0]
            scope_selection_arg = rest[1]
            output_file = rest[2] if len(rest) >= 3 else "engagement_letter_output.docx"

        class _Args:
            pass

        args = _Args()
        args.template = template_file
        args.scope_library = scope_library_file
        args.industry = industry  # may be inferred from variables below
        args.variables = variables_arg
        args.scope_selection = scope_selection_arg
        args.output = output_file

    # Resolve critical file paths up front for clear, actionable failures.
    args.template = _resolve_existing_read_path(
        args.template,
        arg_name="--template",
        purpose="template .docx",
        required=True,
    )
    args.scope_library = _resolve_existing_read_path(
        args.scope_library,
        arg_name="--scope-library",
        purpose="scope library JSON",
        required=True,
    )

    # Normalize output path and ensure parent directory exists.
    try:
        output_path = Path(args.output).expanduser()
        parent = output_path.parent
        if str(parent) not in ("", "."):
            parent.mkdir(parents=True, exist_ok=True)
        args.output = str(output_path)
    except Exception as exc:
        raise SystemExit(f"Invalid --output path: {args.output} ({exc})")

    # Parse variables
    variables = _load_json_arg(args.variables, arg_name="--variables", required=True)

    if getattr(args, "industry", None) in (None, ""):
        inferred = variables.get("INDUSTRY") or variables.get("industry")
        if not inferred:
            raise SystemExit("Missing industry: provide --industry or include INDUSTRY in variables.")
        args.industry = inferred

    scope_selection = None
    if args.scope_selection:
        scope_selection = _load_json_arg(
            args.scope_selection,
            arg_name="--scope-selection",
            required=False,
        )

    result = generate_engagement_letter(
        template_file=args.template,
        scope_library_file=args.scope_library,
        industry=args.industry,
        variables=variables,
        scope_selection=scope_selection,
        output_file=args.output,
    )

    print("\n=== Engagement Letter Generation Summary ===\n")
    for step in result["steps"]:
        print(f"  {step}")
    print(f"\n  Success: {result['success']}")
    if result.get("remaining_placeholders"):
        print(f"\n  Remaining: {result['remaining_placeholders']}")
    print()


if __name__ == "__main__":
    main()
