"""
Engagement Letter Generator (self-contained)

Self-contained script for ChatGPT Code Interpreter. Generates a completed
engagement letter by:
  1. Loading a .docx template
  2. Removing guidance blocks ({{GUIDANCE_NN}} paragraphs)
  3. Replacing the FDD scope block with industry-specific content
  4. Replacing all field/choice placeholders
  5. Validating no {{...}} tokens remain
  6. Saving the filled .docx

Usage (Code Interpreter):
    variables = {"CLIENT_LEGAL_NAME": "Acme Inc.", ...}
    scope_selection = {
        # Optional: omit unchecked top-level scope items (default: include all)
        "excluded_top_level_ids": [],          # e.g. ["scope.002"]
    }
    result = generate_engagement_letter(
        template_file="buyside-engagement-letter.docx",
        scope_library_file="fdd_scope_library.v2.json",
        industry="healthcare",
        variables=variables,
        scope_selection=scope_selection,
        output_file="Engagement_Letter_Acme_Inc.docx"
    )

Usage (CLI):
    python3 el-generate.py \\
        --template buyside-engagement-letter.docx \\
        --scope-library fdd_scope_library.v2.json \\
        --industry healthcare \\
        --variables '{"CLIENT_LEGAL_NAME": "Acme Inc.", ...}' \\
        --scope-selection '{"excluded_top_level_ids":["scope.002"]}' \\
        --output output.docx
"""

import json
import re
import sys
from datetime import date
from dataclasses import dataclass, field
from copy import deepcopy
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Emu
from docx.text.paragraph import Paragraph

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")
GUIDANCE_RE = re.compile(r"\{\{GUIDANCE_\d+\}\}")
GUIDANCE_FREEFORM_RE = re.compile(r"(?i)\{\{\s*GUIDANCE:")

# Placeholders that must be explicitly provided by the user (no inference).
# If present in the selected template and missing/blank at generation time, we refuse.
MUST_ASK_KEYS: set[str] = set()

# Indentation (EMU) matching engagement letter templates
_INDENT_BULLET = Emu(457200)     # 0.50 in — top-level bullet
_INDENT_SUBBULLET = Emu(685800)  # 0.75 in — sub-level bullet
_INDENT_SUBSUBBULLET = Emu(914400)  # 1.00 in — sub-sub-level bullet (rare)

# All guidance placeholder keys (deleted, not replaced)
GUIDANCE_KEYS = {
    f"{{{{GUIDANCE_{str(i).zfill(2)}}}}}" for i in range(1, 19)
}

_REVIEW_PREFIX = "[REVIEW:"

# Block markers (standalone paragraphs) used for deterministic section removal.
BLOCK_INDEPENDENCE_START = "{{BLOCK_INDEPENDENCE_START}}"
BLOCK_INDEPENDENCE_END = "{{BLOCK_INDEPENDENCE_END}}"

# Placeholders where an explicit empty string is a valid default (omits a clause).
ALLOW_EMPTY_KEYS = {
    "EXISTING_SERVICES_CAVEAT",
}

# Numeric money-like fields where we normalize formatting (commas) and support
# legacy "in thousands" inputs (e.g., "80" -> "80,000").
MONEY_KEYS = {
    "FEE_FDD_LOW",
    "FEE_FDD_HIGH",
    "FEE_TAX_LOW",
    "FEE_TAX_HIGH",
}

_MONEY_INT_RE = re.compile(r"^\s*\$?\s*([0-9][0-9,]*)\s*$")


def _is_blank(value) -> bool:
    return value is None or str(value).strip() == ""


def _set_if_blank(out: dict, key: str, value) -> None:
    if _is_blank(out.get(key)):
        out[key] = value


def _format_today(d: Optional[date] = None) -> str:
    d = d or date.today()
    # Platform-independent Month D, YYYY (no leading zero on day).
    return f"{d:%B} {d.day}, {d.year}"


_LEGAL_SUFFIX_RE = re.compile(
    r"""(?ix)
    [,\s]+
    (?:inc\.?|incorporated|ltd\.?|limited|llc|l\.l\.c\.|llp|l\.l\.p\.|lp|l\.p\.|
       corp\.?|corporation|co\.?|company|plc|gmbh|s\.a\.|s\.r\.l\.|sarl|ag)
    \.?$
    """
)


def _client_short_name(client_legal_name: str) -> str:
    s = (client_legal_name or "").strip()
    # Iteratively strip common legal suffixes.
    prev = None
    while s and s != prev:
        prev = s
        s = _LEGAL_SUFFIX_RE.sub("", s).strip().rstrip(",")
    return s or (client_legal_name or "").strip()


def _normalize_money_value(value) -> str:
    """
    Normalize a money-like input to a plain integer string with commas, without
    any currency symbol.

    Back-compat: if the user provides a small integer (< 1000), assume legacy
    "in thousands" and multiply by 1,000.
    """
    if value is None:
        return ""
    s = str(value).strip()
    m = _MONEY_INT_RE.match(s)
    if not m:
        return s
    digits = m.group(1).replace(",", "")
    try:
        n = int(digits)
    except ValueError:
        return s
    if n < 1000:
        n *= 1000
    return f"{n:,}"


def _infer_template_type(template_file: str) -> Optional[str]:
    name = Path(template_file).name.lower()
    if "buyside" in name:
        return "buyside"
    if "sellside" in name:
        return "sellside"
    return None


def _extract_placeholders(doc: Document) -> set[str]:
    """
    Extract placeholder keys (without braces) from the current document state.
    Uses python-docx `.text` so split-across-runs placeholders are recognized.
    """
    keys: set[str] = set()

    def _scan_text(text: str) -> None:
        for token in PLACEHOLDER_RE.findall(text or ""):
            if token in GUIDANCE_KEYS:
                continue
            inner = token[2:-2].strip()
            if inner:
                keys.add(inner)

    for para in doc.paragraphs:
        _scan_text(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan_text(para.text)
                for nested_table in cell.tables:
                    for nrow in nested_table.rows:
                        for ncell in nrow.cells:
                            for para in ncell.paragraphs:
                                _scan_text(para.text)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _scan_text(para.text)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _scan_text(para.text)
        for para in section.footer.paragraphs:
            _scan_text(para.text)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _scan_text(para.text)

    return keys


def _default_review(key: str) -> str:
    return f"[REVIEW: Provide {key}]"


def _has_usable_value(out: dict, key: str) -> bool:
    """
    Return True if the dict has an acceptable value for a placeholder key.
    For some keys, empty string is acceptable to intentionally omit a clause.
    """
    if key not in out:
        return False
    if key in ALLOW_EMPTY_KEYS:
        return out.get(key) is not None
    return not _is_blank(out.get(key))


def derive_and_default_variables(
    *,
    template_file: str,
    variables: dict,
    placeholder_keys: set[str],
) -> dict:
    """
    Apply deterministic inference + defaults before gating/replacement.
    Never overrides a non-blank user-provided value.
    """
    out = dict(variables or {})
    template_type = _infer_template_type(template_file)

    # Dates
    _set_if_blank(out, "LETTER_DATE", _format_today())
    _set_if_blank(out, "DATE_COMMENCE", "Immediately")
    _set_if_blank(out, "DATE_DRAFT_DELIVERY", "TBD")
    _set_if_blank(out, "FAL_DATE", "TBD")
    _set_if_blank(out, "RELEASE_DATE", "TBD")
    if template_type == "sellside":
        _set_if_blank(out, "FAL_LETTER_DATE", out.get("FAL_DATE"))

    # Signing / names
    _set_if_blank(out, "CLIENT_LEGAL_NAME_FULL", out.get("CLIENT_LEGAL_NAME"))
    _set_if_blank(out, "CLIENT_LEGAL_NAME_ACCEPTANCE", out.get("CLIENT_LEGAL_NAME"))
    if _is_blank(out.get("CLIENT_NAME_SHORT")) and not _is_blank(out.get("CLIENT_LEGAL_NAME")):
        out["CLIENT_NAME_SHORT"] = _client_short_name(str(out.get("CLIENT_LEGAL_NAME")))

    _set_if_blank(out, "CLIENT_CONTACT_NAME_TITLE", out.get("CLIENT_CONTACT_NAME"))
    _set_if_blank(out, "EL_SIGNATORY_NAME", out.get("CLIENT_CONTACT_NAME"))
    _set_if_blank(out, "BILLING_ENTITY_NAME", out.get("CLIENT_LEGAL_NAME"))
    _set_if_blank(out, "CLIENT_COUNTRY_2", out.get("CLIENT_COUNTRY"))

    # Signing partners (often same as lead partner)
    _set_if_blank(out, "SIGNING_PARTNER_NAME", out.get("TEAM_PARTNER_NAME"))
    _set_if_blank(out, "SIGNING_PARTNER_NAME_2", out.get("TEAM_PARTNER_NAME"))

    # Project header
    if _is_blank(out.get("PROJECT_NAME_HEADER")) and not _is_blank(out.get("PROJECT_CODE_NAME")):
        code = str(out.get("PROJECT_CODE_NAME")).strip()
        out["PROJECT_NAME_HEADER"] = code if code.lower().startswith("project ") else f"Project {code}"

    # Deal type variants (buyside)
    _set_if_blank(out, "CHOICE_DEAL_TYPE_2", out.get("CHOICE_DEAL_TYPE"))
    if _is_blank(out.get("CHOICE_DEAL_TYPE_3")) and not _is_blank(out.get("CHOICE_DEAL_TYPE")):
        dt = str(out.get("CHOICE_DEAL_TYPE")).strip().lower()
        out["CHOICE_DEAL_TYPE_3"] = "acquisition" if dt == "acquisition of" else out.get("CHOICE_DEAL_TYPE")

    # Target descriptions (used in various header/appendix contexts)
    target_legal = out.get("TARGET_LEGAL_NAME_FULL") or out.get("TARGET_LEGAL_NAME")
    _set_if_blank(out, "TARGET_DESCRIPTION", target_legal)
    _set_if_blank(out, "TARGET_DESCRIPTION_DETAIL", target_legal)

    # Report format default
    _set_if_blank(out, "CHOICE_REPORT_FORMAT", "a PDF written report and an Excel data book")

    # Fields that should remain visibly editable if missing (avoid silent blanks)
    _set_if_blank(out, "EXISTING_SERVICES_DESCRIPTION", "[REVIEW: Insert existing services disclosure]")
    _set_if_blank(out, "EXISTING_SERVICES_CAVEAT", "")
    _set_if_blank(out, "CONFIDENTIALITY_PROCEDURES", "[REVIEW: Insert confidentiality procedures]")
    _set_if_blank(out, "INDEMNITY_EXCLUSION", "[REVIEW: Confirm indemnity exclusion (or leave blank)]")
    _set_if_blank(out, "CLOSING_SENTIMENT", "[REVIEW: Insert closing sentiment]")

    # Fee formatting (full dollars). Also supports legacy "in thousands" inputs.
    for key in MONEY_KEYS:
        if not _is_blank(out.get(key)):
            out[key] = _normalize_money_value(out.get(key))

    # Ensure every placeholder present in the template has a value (never leave raw {{...}}).
    for key in placeholder_keys:
        if not _has_usable_value(out, key):
            out[key] = "" if key in ALLOW_EMPTY_KEYS else _default_review(key)

    return out


# ---------------------------------------------------------------------------
# Optional section removal (marker-based, with heading fallback)
# ---------------------------------------------------------------------------

def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element  # type: ignore[attr-defined]
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _normalize_yes_no(value) -> str:
    s = ("" if value is None else str(value)).strip().lower()
    if s in {"y", "yes", "true", "1"}:
        return "yes"
    if s in {"n", "no", "false", "0"}:
        return "no"
    return ""


def _find_paragraph_index(paragraphs: list[Paragraph], needle: str, start: int = 0) -> int:
    for i in range(start, len(paragraphs)):
        if needle in (paragraphs[i].text or ""):
            return i
    return -1


def apply_independence_section_policy(doc: Document, variables: dict, summary_steps: list[str]) -> None:
    """
    Remove Independence Considerations when it does not apply.

    Preferred: marker-based deletion between BLOCK_INDEPENDENCE_START/END paragraphs.
    Fallback: delete from 'Independence Considerations' heading up to (but not including)
    'Management’s Responsibilities'.
    """
    applies = _normalize_yes_no(variables.get("CHOICE_INDEPENDENCE_APPLIES")) or "yes"

    # Preferred: markers (standalone paragraphs)
    paragraphs = list(doc.paragraphs)
    start_idx = _find_paragraph_index(paragraphs, BLOCK_INDEPENDENCE_START)
    end_idx = _find_paragraph_index(paragraphs, BLOCK_INDEPENDENCE_END, start=max(0, start_idx))

    if start_idx != -1 and end_idx != -1 and end_idx >= start_idx:
        if applies == "no":
            for i in range(end_idx, start_idx - 1, -1):
                _delete_paragraph(paragraphs[i])
            summary_steps.append("Independence Considerations: removed (CHOICE_INDEPENDENCE_APPLIES=no)")
            return

        # applies == "yes": remove marker paragraphs only
        _delete_paragraph(paragraphs[end_idx])
        _delete_paragraph(paragraphs[start_idx])
        summary_steps.append("Independence Considerations: kept (removed block markers)")
        return

    # Fallback: heading-based deletion (only if asked to remove)
    if applies != "no":
        return

    indep_heading = "Independence Considerations"
    mgmt_heading = "Management’s Responsibilities"
    ih = _find_paragraph_index(paragraphs, indep_heading)
    if ih == -1:
        return
    mh = _find_paragraph_index(paragraphs, mgmt_heading, start=ih + 1)
    if mh == -1:
        return

    for i in range(mh - 1, ih - 1, -1):
        _delete_paragraph(paragraphs[i])
    summary_steps.append("Independence Considerations: removed (heading fallback)")


# ---------------------------------------------------------------------------
# Placeholder replacement (run-safe, handles split-across-runs)
# ---------------------------------------------------------------------------

def _replace_in_runs(runs, old: str, new: str) -> int:
    """Replace a placeholder that may span multiple runs. Returns replacement count."""
    count = 0
    # Fast path: single-run match
    for run in runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            count += 1
    if count > 0:
        return count

    # Slow path: placeholder split across runs
    while True:
        full_text = "".join(run.text for run in runs)
        idx = full_text.find(old)
        if idx == -1:
            break
        end_idx = idx + len(old)
        cursor = 0
        inserted = False

        for run in runs:
            text = run.text
            run_end = cursor + len(text)

            if run_end <= idx or cursor >= end_idx:
                cursor = run_end
                continue

            before = text[: max(0, idx - cursor)] if idx > cursor else ""
            after = text[end_idx - cursor:] if run_end > end_idx else ""

            if not inserted:
                run.text = before + str(new) + after
                inserted = True
            else:
                run.text = after

            cursor = run_end

        count += 1

    return count


def _replace_in_paragraph(para, variables: dict) -> None:
    if not para.runs:
        return
    _replace_in_runs(para.runs, variables=None)  # type: ignore


def replace_placeholders(doc: Document, variables: dict) -> None:
    """Replace all {{KEY}} placeholders throughout the document."""
    # Build replacement map: {{KEY}} -> value
    repl = {}
    for key, value in variables.items():
        token = f"{{{{{key}}}}}" if not key.startswith("{{") else key
        repl[token] = str(value) if value is not None else ""

    def _process_paragraph(para):
        if not para.runs:
            return
        for old, new in repl.items():
            _replace_in_runs(para.runs, old, new)

    # Body paragraphs
    for para in doc.paragraphs:
        _process_paragraph(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para)
                for nested_table in cell.tables:
                    for nrow in nested_table.rows:
                        for ncell in nrow.cells:
                            for para in ncell.paragraphs:
                                _process_paragraph(para)

    # Headers/footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _process_paragraph(para)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_paragraph(para)
        for para in section.footer.paragraphs:
            _process_paragraph(para)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _process_paragraph(para)


# ---------------------------------------------------------------------------
# Guidance block removal
# ---------------------------------------------------------------------------

def remove_guidance_blocks(doc: Document) -> int:
    """
    Remove all internal guidance paragraphs from the document.

    Supports both:
    - Numeric placeholders: {{GUIDANCE_NN}}
    - Freeform guidance blocks: {{GUIDANCE: ...}} or [GUIDANCE: ...]
    Returns count of paragraphs removed.
    """
    removed = 0

    def _is_guidance(text: str) -> bool:
        return bool(GUIDANCE_RE.search(text) or GUIDANCE_FREEFORM_RE.search(text))

    def _collect_paragraphs_from_table(table, out: list[Paragraph]) -> None:
        for row in table.rows:
            for cell in row.cells:
                out.extend(cell.paragraphs)
                for nested_table in cell.tables:
                    _collect_paragraphs_from_table(nested_table, out)

    paragraphs: list[Paragraph] = []

    # Body paragraphs
    paragraphs.extend(doc.paragraphs)

    # Tables (body + nested)
    for table in doc.tables:
        _collect_paragraphs_from_table(table, paragraphs)

    # Headers/footers (paragraphs + tables)
    for section in doc.sections:
        paragraphs.extend(section.header.paragraphs)
        for table in section.header.tables:
            _collect_paragraphs_from_table(table, paragraphs)

        paragraphs.extend(section.footer.paragraphs)
        for table in section.footer.tables:
            _collect_paragraphs_from_table(table, paragraphs)

    to_remove: list[Paragraph] = []
    seen = set()
    for para in paragraphs:
        element = getattr(para, "_element", None)
        if element is None:
            continue
        element_id = id(element)
        if element_id in seen:
            continue
        seen.add(element_id)
        if _is_guidance(para.text or ""):
            to_remove.append(para)

    for para in to_remove:
        _delete_paragraph(para)
        removed += 1

    return removed


# ---------------------------------------------------------------------------
# FDD Scope block replacement
# ---------------------------------------------------------------------------

_DIRECTIVE_VERBS = {
    "obtain",
    "read",
    "review",
    "understand",
    "analyze",
    "perform",
    "consider",
    "identify",
    "determine",
    "develop",
    "discuss",
    "evaluate",
    "compare",
    "summarize",
    "comment",
    "meet",
    "inquire",
    "assess",
    "bridge",
    "reconcile",
    "gain",
    "propose",
    "segment",
    "calculate",
    "collect",
    "confirm",
    "test",
    "prepare",
    "provide",
}


def _first_word(text: str) -> str:
    m = re.match(r"[\W_]*(\w+)", text.strip())
    return (m.group(1) if m else "").lower()


def _looks_like_new_directive(text: str) -> bool:
    w = _first_word(text)
    return w in _DIRECTIVE_VERBS and len(text.strip().split()) >= 3


@dataclass
class _ScopeNode:
    text: str
    children: list["_ScopeNode"] = field(default_factory=list)
    top_level_id: Optional[str] = None


def _is_v2_bullet_list(bullets: list) -> bool:
    """Return True if bullets appear to be v2 objects with text/children."""
    if not isinstance(bullets, list) or not bullets:
        return False
    sample = bullets[0]
    return isinstance(sample, dict) and isinstance(sample.get("text"), str)


def _parse_scope_bullets(bullets: list) -> list[_ScopeNode]:
    """
    Convert a flat list of scope bullet strings into a simple parent/child tree.

    Heuristic:
    - Lines ending with ':' become parents; following lines attach beneath them.
    - A “new directive” line (e.g., starts with obtain/analyze/perform/consider/...) ends the current nesting.
    """
    root: list[_ScopeNode] = []
    stack: list[_ScopeNode] = []

    def _add(text: str) -> _ScopeNode:
        node = _ScopeNode(text=text)
        if stack:
            stack[-1].children.append(node)
        else:
            root.append(node)
        return node

    for bullet in bullets:
        if not isinstance(bullet, str):
            continue
        text = bullet.strip()
        if not text:
            continue

        if text.endswith(":"):
            if _looks_like_new_directive(text):
                stack.clear()
            node = _add(text)
            stack.append(node)
            continue

        if stack and _looks_like_new_directive(text):
            stack.clear()
        _add(text)

    return root


def _scope_nodes_from_v2(*, bullets: list) -> list[_ScopeNode]:
    """
    Build scope nodes from v2 bullet objects.

    v2 node shape:
      {"id": "scope.001", "text": "...", "children": [ ... ]}
    """
    def build(node: dict) -> _ScopeNode:
        text = node.get("text", "") if isinstance(node, dict) else ""
        n = _ScopeNode(text=text, children=[], top_level_id=node.get("id") if isinstance(node, dict) else None)
        for child in (node.get("children") or []) if isinstance(node, dict) else []:
            if isinstance(child, dict):
                n.children.append(build(child))
        return n

    out: list[_ScopeNode] = []
    for node in bullets:
        if isinstance(node, dict) and isinstance(node.get("text"), str):
            out.append(build(node))
    return out


def _scope_nodes_from_schema(*, bullets: list, schema_nodes: list) -> list[_ScopeNode]:
    """
    Build scope nodes from explicit schema (indices into the bullets list).

    Schema node shape:
      {"i": <index>, "id": "scope.xxxxxxxx", "children": [ ... ]}
    """
    def build(node: dict) -> _ScopeNode:
        i = int(node["i"])
        text = bullets[i] if 0 <= i < len(bullets) and isinstance(bullets[i], str) else ""
        n = _ScopeNode(text=text, children=[], top_level_id=node.get("id"))
        for child in (node.get("children") or []):
            if isinstance(child, dict):
                n.children.append(build(child))
        return n

    out: list[_ScopeNode] = []
    for node in schema_nodes:
        if isinstance(node, dict):
            out.append(build(node))
    return out


def _assign_top_level_ids(nodes: list[_ScopeNode], start: int) -> int:
    """
    Assign stable ids to top-level scope nodes in document insertion order.

    IDs are global (not section-specific) to avoid exposing internal section keys in Canvas.
    """
    n = start
    for node in nodes:
        if not node.top_level_id:
            node.top_level_id = f"scope.{n:03d}"
        n += 1
    return n


def _indent_for_depth(depth: int) -> Emu:
    if depth <= 0:
        return _INDENT_BULLET
    if depth == 1:
        return _INDENT_SUBBULLET
    return _INDENT_SUBSUBBULLET


def load_scope_library(path: str) -> dict:
    """Load the FDD scope library JSON."""
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def replace_fdd_scope_block(
    doc: Document,
    industry: str,
    scope_library: dict,
    scope_selection: Optional[dict] = None,
    scope_heading_search: str = "FINANCIAL DUE DILIGENCE",
    end_boundary_search: str = "These Terms and Conditions",
) -> dict:
    """
    Replace the sample FDD scope block with industry-specific content.

    Returns summary dict with counts.
    """
    paras = list(doc.paragraphs)

    # Find scope boundaries
    start_idx = None
    end_idx = None
    for i, p in enumerate(paras):
        if start_idx is None and scope_heading_search in p.text:
            start_idx = i
        elif start_idx is not None and p.text.strip().startswith(end_boundary_search):
            end_idx = i
            break

    if start_idx is None:
        return {"error": f"Scope heading '{scope_heading_search}' not found", "sections_inserted": 0, "bullets_inserted": 0}
    if end_idx is None:
        return {"error": f"End boundary '{end_boundary_search}' not found", "sections_inserted": 0, "bullets_inserted": 0}

    # Capture formatting templates
    bold_template = None
    bullet_template = None
    for p in paras[start_idx + 1: end_idx]:
        if not p.text.strip():
            continue
        if bold_template is None and p.runs and any(r.bold for r in p.runs):
            bold_template = p
        if bullet_template is None and p.runs and not any(r.bold for r in p.runs):
            bullet_template = p

    if bold_template is None:
        for p in paras[start_idx + 1: end_idx]:
            if p.text.strip():
                bold_template = p
                break
    if bullet_template is None:
        bullet_template = bold_template

    if bold_template is None:
        return {"error": "No formatting template found in scope block", "sections_inserted": 0, "bullets_inserted": 0}

    # Delete scope content
    to_remove = paras[start_idx + 1: end_idx]
    for p in to_remove:
        p._p.getparent().remove(p._p)

    # Build and insert replacement content
    common = scope_library.get("common_skeleton", [])
    modules = scope_library.get("industry_modules", {}).get(industry, {})
    common_headings = {s["normalized_heading"] for s in common}

    extra_sections = [
        (heading, bullets)
        for heading, bullets in modules.items()
        if heading not in common_headings
    ]

    excluded_top_level_ids: set[str] = set()
    if isinstance(scope_selection, dict):
        excluded_top_level_ids = set(
            scope_selection.get("excluded_top_level_ids", [])
            or scope_selection.get("exclude_top_level_ids", [])
            or scope_selection.get("excluded_top_level_bullet_ids", [])
            or []
        )

    cursor = paras[start_idx]
    section_count = 0
    bullet_count = 0
    next_scope_id = 1

    scope_schema = scope_library.get("scope_schema") if isinstance(scope_library, dict) else None
    schema_nesting = scope_schema.get("nesting") if isinstance(scope_schema, dict) else None
    common_schema = schema_nesting.get("common_skeleton") if isinstance(schema_nesting, dict) else None
    industry_schema = schema_nesting.get("industry_modules") if isinstance(schema_nesting, dict) else None

    def _insert_after(cursor_para, text, bold=False, indent=None):
        template = bold_template if bold else bullet_template
        new_p = deepcopy(template._p)
        cursor_para._p.addnext(new_p)
        new_para = Paragraph(new_p, cursor_para._parent)

        if new_para.runs:
            new_para.runs[0].text = text
            for run in new_para.runs[1:]:
                run.text = ""
            for run in new_para.runs:
                run.bold = bold
        else:
            run = new_para.add_run(text)
            run.bold = bold

        pf = new_para.paragraph_format
        pf.left_indent = indent

        try:
            new_para.style = doc.styles["Body Text"]
        except KeyError:
            pass

        return new_para

    def _insert_nodes(section_key: str, nodes: list[_ScopeNode]) -> None:
        nonlocal cursor, bullet_count
        for node in nodes:
            if node.top_level_id and node.top_level_id in excluded_top_level_ids:
                continue

            cursor = _insert_after(cursor, node.text, indent=_indent_for_depth(0))
            bullet_count += 1

            def _insert_children(children: list[_ScopeNode], depth: int) -> None:
                nonlocal cursor, bullet_count
                for child in children:
                    cursor = _insert_after(cursor, child.text, indent=_indent_for_depth(depth))
                    bullet_count += 1
                    if child.children:
                        _insert_children(child.children, depth + 1)

            if node.children:
                _insert_children(node.children, 1)

    # Common skeleton sections
    for section in common:
        heading = section["heading"]
        norm = section["normalized_heading"]
        default_bullets = section.get("default_bullets", [])
        industry_bullets = modules.get(norm, [])

        nodes_default = None
        if _is_v2_bullet_list(default_bullets):
            nodes_default = _scope_nodes_from_v2(bullets=default_bullets)
        elif isinstance(common_schema, dict):
            sec = common_schema.get(norm)
            if isinstance(sec, dict):
                sn = sec.get("default_bullets")
                if isinstance(sn, list):
                    nodes_default = _scope_nodes_from_schema(bullets=default_bullets, schema_nodes=sn)

        if nodes_default is None:
            nodes_default = _parse_scope_bullets(default_bullets)
        next_scope_id = _assign_top_level_ids(nodes_default, next_scope_id)

        # Industry bullets that belong under this common section
        nodes_industry = None
        if _is_v2_bullet_list(industry_bullets):
            nodes_industry = _scope_nodes_from_v2(bullets=industry_bullets)
        elif isinstance(industry_schema, dict):
            ind = industry_schema.get(industry)
            if isinstance(ind, dict):
                sn = ind.get(norm)
                if isinstance(sn, list):
                    nodes_industry = _scope_nodes_from_schema(bullets=industry_bullets, schema_nodes=sn)

        if nodes_industry is None:
            nodes_industry = _parse_scope_bullets(industry_bullets)
        next_scope_id = _assign_top_level_ids(nodes_industry, next_scope_id)

        nodes = nodes_default + nodes_industry

        if not any(n.top_level_id not in excluded_top_level_ids for n in nodes):
            continue

        cursor = _insert_after(cursor, heading, bold=True)
        section_count += 1
        _insert_nodes(norm, nodes)

    # Extra industry-only sections
    for heading_key, bullets in extra_sections:
        display_heading = heading_key.replace("_", " ").title()
        nodes = None
        if _is_v2_bullet_list(bullets):
            nodes = _scope_nodes_from_v2(bullets=bullets)
        elif isinstance(industry_schema, dict):
            ind = industry_schema.get(industry)
            if isinstance(ind, dict):
                sn = ind.get(heading_key)
                if isinstance(sn, list):
                    nodes = _scope_nodes_from_schema(bullets=bullets, schema_nodes=sn)

        if nodes is None:
            nodes = _parse_scope_bullets(bullets)
        next_scope_id = _assign_top_level_ids(nodes, next_scope_id)

        if not any(n.top_level_id not in excluded_top_level_ids for n in nodes):
            continue

        cursor = _insert_after(cursor, display_heading, bold=True)
        section_count += 1
        _insert_nodes(heading_key, nodes)

    return {
        "sections_inserted": section_count,
        "bullets_inserted": bullet_count,
        "industry": industry,
    }


# ---------------------------------------------------------------------------
# Validation
# ---------------------------------------------------------------------------

def validate_no_placeholders(doc: Document) -> list:
    """Return list of any remaining {{...}} placeholders."""
    remaining = []

    def _scan(para):
        matches = PLACEHOLDER_RE.findall(para.text)
        remaining.extend(matches)

    for para in doc.paragraphs:
        _scan(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan(para)
    for section in doc.sections:
        for para in section.header.paragraphs:
            _scan(para)
        for para in section.footer.paragraphs:
            _scan(para)

    return sorted(set(remaining))


# ---------------------------------------------------------------------------
# Main generation pipeline
# ---------------------------------------------------------------------------

def generate_engagement_letter(
    template_file: str,
    scope_library_file: str,
    industry: str,
    variables: dict,
    scope_selection: Optional[dict] = None,
    output_file: str = "engagement_letter_output.docx",
) -> dict:
    """
    Generate a completed engagement letter.

    Args:
        template_file: Path to the .docx template.
        scope_library_file: Path to the scope library JSON (e.g. fdd_scope_library.v2.json).
        industry: Industry key (e.g. "healthcare").
        variables: Dict of KEY -> value (without {{ }} wrappers).
        output_file: Output filename.

    Returns:
        Summary dict with generation results.
    """
    summary = {
        "template": template_file,
        "industry": industry,
        "output": output_file,
        "steps": [],
        "success": False,
    }

    # Step 1: Load template
    doc = Document(template_file)
    summary["steps"].append(f"Loaded template: {template_file}")

    # Step 1.5: Optional section removal (before placeholder extraction)
    apply_independence_section_policy(doc, variables or {}, summary["steps"])

    # Step 2: Remove guidance blocks
    guidance_removed = remove_guidance_blocks(doc)
    summary["steps"].append(f"Removed {guidance_removed} guidance block paragraphs")

    # Step 3: Replace FDD scope block
    scope_library = load_scope_library(scope_library_file)
    scope_result = replace_fdd_scope_block(
        doc,
        industry=industry,
        scope_library=scope_library,
        scope_selection=scope_selection,
    )
    if "error" in scope_result:
        summary["steps"].append(f"Scope replacement warning: {scope_result['error']}")
    else:
        summary["steps"].append(
            f"Inserted FDD scope: {scope_result['sections_inserted']} sections, "
            f"{scope_result['bullets_inserted']} bullets"
        )

    # Step 4: Derive defaults + ensure template placeholder coverage (pre-gate)
    placeholder_keys = _extract_placeholders(doc)
    variables = derive_and_default_variables(
        template_file=template_file,
        variables=variables,
        placeholder_keys=placeholder_keys,
    )

    missing_must_ask = [
        k for k in sorted(MUST_ASK_KEYS & placeholder_keys)
        if _is_blank(variables.get(k)) or str(variables.get(k)).startswith(_REVIEW_PREFIX)
    ]
    if missing_must_ask:
        summary["steps"].append(f"ERROR: Missing required user-provided fields: {missing_must_ask}")
        summary["missing_required_fields"] = missing_must_ask
        summary["success"] = False
        return summary

    # Conditional must-ask: SEC/audit status is required only if independence section is present.
    if "CHOICE_SEC_STATUS" in placeholder_keys:
        applies = _normalize_yes_no(variables.get("CHOICE_INDEPENDENCE_APPLIES")) or "yes"
        if applies == "yes" and (
            _is_blank(variables.get("CHOICE_SEC_STATUS"))
            or str(variables.get("CHOICE_SEC_STATUS")).startswith(_REVIEW_PREFIX)
        ):
            summary["steps"].append("ERROR: Missing required user-provided fields: ['CHOICE_SEC_STATUS']")
            summary["missing_required_fields"] = ["CHOICE_SEC_STATUS"]
            summary["success"] = False
            return summary

    summary["steps"].append(f"Derived defaults + filled template placeholders ({len(placeholder_keys)} keys)")

    # Step 5: Replace all field/choice placeholders
    replace_placeholders(doc, variables)
    summary["steps"].append(f"Replaced placeholders ({len(variables)} variables)")

    # Step 6: Validate
    remaining = validate_no_placeholders(doc)
    if remaining:
        summary["steps"].append(f"WARNING: {len(remaining)} unreplaced placeholders: {remaining}")
        summary["remaining_placeholders"] = remaining
    else:
        summary["steps"].append("Validation passed: no remaining placeholders")

    # Step 7: Save
    doc.save(output_file)
    summary["steps"].append(f"Saved: {output_file}")
    summary["success"] = len(remaining) == 0

    return summary


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    import argparse

    # Back-compat: Code Interpreter sometimes invokes this script with positional
    # args instead of flags. Support both:
    #   python el-generate.py --template ... --scope-library ... --industry ... --variables ... --scope-selection ... --output ...
    #   python el-generate.py TEMPLATE_FILE SCOPE_LIBRARY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]
    #   python el-generate.py TEMPLATE_FILE SCOPE_LIBRARY INDUSTRY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]
    raw_argv = sys.argv[1:]
    has_flag = any(a.startswith("--") for a in raw_argv)

    if has_flag:
        parser = argparse.ArgumentParser(description="Generate an engagement letter")
        parser.add_argument("--template", required=True, help="Path to .docx template")
        parser.add_argument("--scope-library", required=True, help="Path to the scope library JSON (e.g. fdd_scope_library.v2.json)")
        parser.add_argument("--industry", required=True, help="Industry key")
        parser.add_argument("--variables", required=True, help="JSON string or path to JSON file with variables")
        parser.add_argument("--scope-selection", help="Optional JSON string or path to JSON file controlling scope inclusion")
        parser.add_argument("--output", default="engagement_letter_output.docx", help="Output filename")
        args = parser.parse_args()
    else:
        if len(raw_argv) < 4:
            raise SystemExit(
                "Usage:\n"
                "  el-generate.py --template ... --scope-library ... --industry ... --variables ... [--scope-selection ...] [--output ...]\n"
                "  el-generate.py TEMPLATE_FILE SCOPE_LIBRARY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]\n"
                "  el-generate.py TEMPLATE_FILE SCOPE_LIBRARY INDUSTRY VARIABLES_JSON SCOPE_SELECTION_JSON [OUTPUT_FILE]"
            )

        template_file = raw_argv[0]
        scope_library_file = raw_argv[1]
        rest = raw_argv[2:]

        def _is_jsonish(s: str) -> bool:
            s2 = (s or "").lstrip()
            return s2.startswith("{") or s2.startswith("[")

        industry = None
        if len(rest) >= 3 and (not _is_jsonish(rest[0])) and (Path(rest[1]).exists() or _is_jsonish(rest[1])):
            # TEMPLATE SCOPE INDUSTRY VARIABLES SCOPE_SELECTION [OUTPUT]
            industry = rest[0]
            variables_arg = rest[1]
            scope_selection_arg = rest[2]
            output_file = rest[3] if len(rest) >= 4 else "engagement_letter_output.docx"
        else:
            # TEMPLATE SCOPE VARIABLES SCOPE_SELECTION [OUTPUT]
            variables_arg = rest[0]
            scope_selection_arg = rest[1]
            output_file = rest[2] if len(rest) >= 3 else "engagement_letter_output.docx"

        class _Args:
            pass

        args = _Args()
        args.template = template_file
        args.scope_library = scope_library_file
        args.industry = industry  # may be inferred from variables below
        args.variables = variables_arg
        args.scope_selection = scope_selection_arg
        args.output = output_file

    # Parse variables
    if Path(args.variables).exists():
        with open(args.variables, "r") as f:
            variables = json.load(f)
    else:
        variables = json.loads(args.variables)

    if getattr(args, "industry", None) in (None, ""):
        inferred = variables.get("INDUSTRY") or variables.get("industry")
        if not inferred:
            raise SystemExit("Missing industry: provide --industry or include INDUSTRY in variables.")
        args.industry = inferred

    scope_selection = None
    if args.scope_selection:
        if Path(args.scope_selection).exists():
            with open(args.scope_selection, "r") as f:
                scope_selection = json.load(f)
        else:
            scope_selection = json.loads(args.scope_selection)

    result = generate_engagement_letter(
        template_file=args.template,
        scope_library_file=args.scope_library,
        industry=args.industry,
        variables=variables,
        scope_selection=scope_selection,
        output_file=args.output,
    )

    print("\n=== Engagement Letter Generation Summary ===\n")
    for step in result["steps"]:
        print(f"  {step}")
    print(f"\n  Success: {result['success']}")
    if result.get("remaining_placeholders"):
        print(f"\n  Remaining: {result['remaining_placeholders']}")
    print()


if __name__ == "__main__":
    main()
