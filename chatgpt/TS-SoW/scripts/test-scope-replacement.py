"""
Test / Demo: FDD Scope Block Replacement

Loads both engagement letter templates, replaces the sample FDD scope block
with industry-specific content from the scope library, and saves output files
for visual verification in Word.

Usage:
    python3 scripts/test-scope-replacement.py

    # Use schema-annotated library copy (keeps the original library untouched):
    python3 scripts/test-scope-replacement.py --scope-library reference/fdd_scope_library.schema_v1_1.json
"""

import sys
from pathlib import Path

# Ensure the project root is on the Python path so we can import reference modules
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document
from reference.helper_functions import load_scope_library, replace_fdd_scope_block
from reference.el_template_schema import TEMPLATES


def run_test(template_key: str, industry: str, output_dir: Path, scope_library_path: Path) -> None:
    """Run scope replacement on one template and save the result."""
    config = TEMPLATES[template_key]
    template_path = PROJECT_ROOT / config["filename"]
    scope_config = config.get("scope_config", {})

    short_name = template_key.replace("_engagement_letter", "")
    output_path = output_dir / f"test-{short_name}-{industry}.docx"

    print(f"\n{'=' * 60}")
    print(f"Template:  {config['display_name']}")
    print(f"Industry:  {industry}")
    print(f"Input:     {template_path}")
    print(f"Output:    {output_path}")
    print(f"{'=' * 60}")

    # Load template
    doc = Document(str(template_path))
    total_paras_before = len(doc.paragraphs)

    # Load scope library
    scope_library = load_scope_library(str(scope_library_path))

    # Replace scope block
    result = replace_fdd_scope_block(
        doc,
        industry=industry,
        scope_library=scope_library,
        scope_heading_search=scope_config.get("scope_heading_search", "FINANCIAL DUE DILIGENCE"),
        end_boundary_search=scope_config.get("end_boundary_search", "These Terms and Conditions"),
    )

    total_paras_after = len(doc.paragraphs)

    # Save
    doc.save(str(output_path))

    # Print summary
    print(f"\nResult summary:")
    print(f"  Sections inserted:          {result['sections_inserted']}")
    print(f"  Bullets inserted:           {result['bullets_inserted']}")
    print(f"  Common skeleton sections:   {result['common_sections']}")
    print(f"  Extra industry sections:    {result['extra_industry_sections']}")
    print(f"  Paragraphs before:          {total_paras_before}")
    print(f"  Paragraphs after:           {total_paras_after}")
    print(f"  Saved to: {output_path}")


def main():
    import argparse

    parser = argparse.ArgumentParser(description="Test FDD scope block replacement")
    parser.add_argument("--industry", default="healthcare", help="Industry key (e.g., healthcare, tech)")
    parser.add_argument(
        "--scope-library",
        default="reference/fdd_scope_library.json",
        help="Path to scope library JSON (can be schema-annotated copy)",
    )
    args = parser.parse_args()

    output_dir = PROJECT_ROOT / "dist"
    output_dir.mkdir(exist_ok=True)

    test_industry = args.industry
    scope_library_path = (PROJECT_ROOT / args.scope_library).resolve()

    # Verify the industry exists in the scope library
    scope_library = load_scope_library(str(scope_library_path))
    available = list(scope_library.get("industry_modules", {}).keys())
    if test_industry not in available:
        print(f"ERROR: Industry '{test_industry}' not found. Available: {available}")
        sys.exit(1)

    print(f"FDD Scope Block Replacement Test")
    print(f"Industry: {test_industry}")
    print(f"Available industries: {available}")

    # Run on both templates
    for template_key in ["buyside_engagement_letter", "sellside_engagement_letter"]:
        try:
            run_test(template_key, test_industry, output_dir, scope_library_path)
        except Exception as e:
            print(f"\nERROR on {template_key}: {e}")
            import traceback
            traceback.print_exc()

    print(f"\n{'=' * 60}")
    print("Done! Open the output files in Word to verify formatting.")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
