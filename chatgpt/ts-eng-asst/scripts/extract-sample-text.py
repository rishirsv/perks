"""
Extract structured text from all SOW sample documents (docx + pdf).

Walks dist/el-samples/, extracts paragraphs with formatting metadata from each
.docx and .pdf, and saves the output to dist/el-samples-text/{industry}/{filename}.json.

Each paragraph includes: text, Word style name, bold flag, indent level, and font size.
This preserves structural hierarchy (headings vs. bullets vs. body) so downstream
agents can categorize content without guessing.

Usage:
    python scripts/extract-sample-text.py
"""

import json
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
SAMPLES_DIR = REPO_ROOT / "dist" / "el-samples"
OUTPUT_DIR = REPO_ROOT / "dist" / "el-samples-text"


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def _get_font_size_pt(para) -> float | None:
    """Get the effective font size in points for a paragraph."""
    # Check run-level first (most specific)
    for run in para.runs:
        if run.font.size is not None:
            return run.font.size.pt
    # Check paragraph style
    if para.style and para.style.font and para.style.font.size:
        return para.style.font.size.pt
    return None


def _get_bold(para) -> bool:
    """Check if the paragraph's first run is bold, or if the style is bold."""
    for run in para.runs:
        if run.text.strip():
            if run.bold is True:
                return True
            break
    if para.style and para.style.font and para.style.font.bold:
        return True
    return False


def _get_indent_level(para) -> int:
    """Infer indent level from left indent or list level."""
    indent = para.paragraph_format.left_indent
    if indent is not None:
        # Convert EMUs to approximate indent level (each level ~360000 EMUs / 0.25in)
        emu = int(indent)
        if emu > 0:
            return max(1, round(emu / 360000))
    # Check for list-style indentation in style name
    style_name = (para.style.name or "").lower()
    if "list" in style_name:
        # "List Bullet 2" -> level 2, "List Bullet" -> level 1
        for part in style_name.split():
            if part.isdigit():
                return int(part)
        return 1
    return 0


def extract_docx(filepath: Path) -> dict:
    """Extract structured paragraph data from a .docx file."""
    from docx import Document

    doc = Document(str(filepath))
    paragraphs: list[dict] = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        paragraphs.append({
            "index": i,
            "text": text,
            "style": para.style.name if para.style else "Normal",
            "bold": _get_bold(para),
            "indent_level": _get_indent_level(para),
            "font_size_pt": _get_font_size_pt(para),
        })

    tables: list[dict] = []
    for ti, table in enumerate(doc.tables):
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({
            "table_index": ti,
            "rows": rows,
        })

    headers_footers: list[dict] = []
    for si, section in enumerate(doc.sections):
        for para in section.header.paragraphs:
            text = para.text.strip()
            if text:
                headers_footers.append({"location": "header", "section": si, "text": text})
        for para in section.footer.paragraphs:
            text = para.text.strip()
            if text:
                headers_footers.append({"location": "footer", "section": si, "text": text})

    return {
        "source_file": filepath.name,
        "format": "docx",
        "paragraphs": paragraphs,
        "tables": tables,
        "headers_footers": headers_footers,
    }


def extract_pdf(filepath: Path) -> dict:
    """Extract text from a .pdf file. Limited structural metadata available."""
    text_pages: list[dict] = []

    try:
        import pdfplumber
        with pdfplumber.open(str(filepath)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if text.strip():
                    text_pages.append({"page": i + 1, "text": text.strip()})
    except ImportError:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(filepath))
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    text_pages.append({"page": i + 1, "text": text.strip()})
        except ImportError:
            text_pages.append({"page": 0, "text": "[ERROR] No PDF library available."})

    # PDFs don't provide paragraph-level style metadata, so we store raw pages
    return {
        "source_file": filepath.name,
        "format": "pdf",
        "pages": text_pages,
        "note": "PDF extraction provides raw text only. No paragraph-level style metadata available.",
    }


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    if not SAMPLES_DIR.exists():
        print(f"ERROR: Samples directory not found: {SAMPLES_DIR}")
        sys.exit(1)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    stats = {"docx": 0, "pdf": 0, "skipped": 0, "errors": 0}
    results: list[str] = []

    for industry_dir in sorted(SAMPLES_DIR.iterdir()):
        if not industry_dir.is_dir():
            continue

        industry_name = industry_dir.name
        out_industry = OUTPUT_DIR / industry_name
        out_industry.mkdir(parents=True, exist_ok=True)

        for filepath in sorted(industry_dir.iterdir()):
            if filepath.name.startswith("~$"):
                stats["skipped"] += 1
                continue

            suffix = filepath.suffix.lower()
            stem = filepath.stem

            if suffix == ".docx":
                try:
                    data = extract_docx(filepath)
                    stats["docx"] += 1
                except Exception as exc:
                    data = {"source_file": filepath.name, "error": str(exc)}
                    stats["errors"] += 1

            elif suffix == ".pdf":
                try:
                    data = extract_pdf(filepath)
                    stats["pdf"] += 1
                except Exception as exc:
                    data = {"source_file": filepath.name, "error": str(exc)}
                    stats["errors"] += 1

            else:
                stats["skipped"] += 1
                continue

            out_file = out_industry / f"{stem}.json"
            out_file.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")

            para_count = len(data.get("paragraphs", data.get("pages", [])))
            results.append(f"  {industry_name}/{stem}.json ({para_count} paragraphs)")

    total = stats["docx"] + stats["pdf"]
    print(f"Extraction complete: {total} files processed")
    print(f"  docx: {stats['docx']}, pdf: {stats['pdf']}, skipped: {stats['skipped']}, errors: {stats['errors']}")
    print(f"  Output: {OUTPUT_DIR}")
    print()
    for r in results:
        print(r)


if __name__ == "__main__":
    main()
