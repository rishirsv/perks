"""
Test / Demo: FDD Scope Block Replacement (uses dist generator logic)

Loads an engagement letter template, replaces the sample FDD scope block with
industry-specific content from the scope library, and saves output files for
visual verification in Word.

Usage:
  python3 scripts/test-scope-replacement.py --industry retail
  python3 scripts/test-scope-replacement.py --industry generic
"""

from __future__ import annotations

import argparse
import importlib.util
import json
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parent.parent


def _load_generator_module():
    mod_path = PROJECT_ROOT / "dist" / "el-generate.py"
    spec = importlib.util.spec_from_file_location("el_generate", mod_path)
    if not spec or not spec.loader:
        raise RuntimeError(f"Failed to load generator module: {mod_path}")
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def _load_scope_library(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def run_test(*, template_name: str, industry: str, output_dir: Path, scope_library: dict) -> None:
    template_path = PROJECT_ROOT / "dist" / template_name
    output_path = output_dir / f"test-{template_path.stem}-{industry}.docx"

    print(f"\n{'=' * 60}")
    print(f"Template:  {template_name}")
    print(f"Industry:  {industry}")
    print(f"Input:     {template_path}")
    print(f"Output:    {output_path}")
    print(f"{'=' * 60}")

    doc = Document(str(template_path))
    total_paras_before = len(doc.paragraphs)

    gen = _load_generator_module()
    result = gen.replace_fdd_scope_block(
        doc,
        industry=industry,
        scope_library=scope_library,
        scope_heading_search="FINANCIAL DUE DILIGENCE",
        end_boundary_search="These Terms and Conditions",
    )

    total_paras_after = len(doc.paragraphs)
    doc.save(str(output_path))

    print("\nResult summary:")
    if "error" in result:
        print(f"  Error: {result['error']}")
        return
    print(f"  Sections inserted:          {result['sections_inserted']}")
    print(f"  Bullets inserted:           {result['bullets_inserted']}")
    print(f"  Common skeleton sections:   {result['common_sections']}")
    print(f"  Extra industry sections:    {result['extra_industry_sections']}")
    print(f"  Paragraphs before:          {total_paras_before}")
    print(f"  Paragraphs after:           {total_paras_after}")
    print(f"  Saved to: {output_path}")


def main() -> None:
    parser = argparse.ArgumentParser(description="Test FDD scope block replacement")
    parser.add_argument("--industry", default="healthcare", help="Industry key (e.g., healthcare, retail, generic)")
    parser.add_argument(
        "--scope-library",
        default="dist/fdd_scope_library.v2.json",
        help="Path to scope library JSON (v2 bundle recommended)",
    )
    parser.add_argument(
        "--out-dir",
        default="docs/scope-library/generated",
        help="Output directory for generated .docx files (non-dist).",
    )
    args = parser.parse_args()

    output_dir = (PROJECT_ROOT / args.out_dir).resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    scope_library_path = (PROJECT_ROOT / args.scope_library).resolve()
    scope_library = _load_scope_library(scope_library_path)
    available = sorted((scope_library.get("industry_modules") or {}).keys())

    print("FDD Scope Block Replacement Test")
    print(f"Industry: {args.industry}")
    print(f"Available industries: {available}")

    for template_name in ["buyside-engagement-letter.docx", "sellside-engagement-letter.docx"]:
        run_test(
            template_name=template_name,
            industry=args.industry,
            output_dir=output_dir,
            scope_library=scope_library,
        )

    print(f"\n{'=' * 60}")
    print("Done! Open the output files in Word to verify formatting.")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()

