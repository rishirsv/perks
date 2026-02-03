"""
Audit Placeholders in Engagement Letter Templates

Extracts every {{...}} pattern from both buyside and sellside templates,
reporting paragraph index, placeholder text, and surrounding context.

Usage:
    python3 scripts/audit-placeholders.py
"""

import re
import sys
from pathlib import Path

PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from docx import Document

PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")


def audit_template(docx_path: Path) -> list[dict]:
    """Extract all placeholders with context from a .docx file."""
    doc = Document(str(docx_path))
    results = []

    # Body paragraphs
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for match in PLACEHOLDER_RE.finditer(text):
            results.append({
                "location": "body",
                "paragraph_index": i,
                "placeholder": match.group(),
                "context": text.strip()[:120],
            })

    # Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = para.text
                    for match in PLACEHOLDER_RE.finditer(text):
                        results.append({
                            "location": f"table[{t_idx}] row[{r_idx}] cell[{c_idx}]",
                            "paragraph_index": p_idx,
                            "placeholder": match.group(),
                            "context": text.strip()[:120],
                        })

    # Headers/footers
    for s_idx, section in enumerate(doc.sections):
        for h_idx, para in enumerate(section.header.paragraphs):
            text = para.text
            for match in PLACEHOLDER_RE.finditer(text):
                results.append({
                    "location": f"header[{s_idx}]",
                    "paragraph_index": h_idx,
                    "placeholder": match.group(),
                    "context": text.strip()[:120],
                })
        for f_idx, para in enumerate(section.footer.paragraphs):
            text = para.text
            for match in PLACEHOLDER_RE.finditer(text):
                results.append({
                    "location": f"footer[{s_idx}]",
                    "paragraph_index": f_idx,
                    "placeholder": match.group(),
                    "context": text.strip()[:120],
                })

    return results


def main():
    templates = {
        "buyside": PROJECT_ROOT / "dist" / "buyside-engagement-letter.docx",
        "sellside": PROJECT_ROOT / "dist" / "sellside-engagement-letter.docx",
    }

    for name, path in templates.items():
        print(f"\n{'=' * 70}")
        print(f"TEMPLATE: {name} ({path.name})")
        print(f"{'=' * 70}\n")

        if not path.exists():
            print(f"  FILE NOT FOUND: {path}")
            continue

        results = audit_template(path)

        # Unique placeholders
        unique = sorted(set(r["placeholder"] for r in results))
        guidance = [p for p in unique if p.upper().startswith("{{GUIDANCE")]
        fields = [p for p in unique if not p.upper().startswith("{{GUIDANCE")]

        print(f"Total occurrences: {len(results)}")
        print(f"Unique placeholders: {len(unique)}")
        print(f"  Field/choice: {len(fields)}")
        print(f"  Guidance:     {len(guidance)}")

        print(f"\n--- All occurrences (in document order) ---\n")
        for r in results:
            print(f"  [{r['location']} ¶{r['paragraph_index']}]")
            print(f"    Placeholder: {r['placeholder']}")
            print(f"    Context:     {r['context']}")
            print()

        print(f"\n--- Unique field/choice placeholders ---\n")
        for p in fields:
            count = sum(1 for r in results if r["placeholder"] == p)
            print(f"  {p}  (x{count})")

        print(f"\n--- Guidance blocks ---\n")
        for p in guidance:
            ctx = next(r["context"] for r in results if r["placeholder"] == p)
            print(f"  {p}")
            print(f"    Context: {ctx}")
            print()


if __name__ == "__main__":
    main()
