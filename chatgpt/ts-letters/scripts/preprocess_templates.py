#!/usr/bin/env python3
"""
Preprocess DOCX templates to standardize placeholder casing.

This script:
1. Reads templates from templates/ folder
2. Standardizes placeholder casing (e.g., {{DATE}} -> {{Date}})
3. Saves processed templates to dist/ folder
4. Generates a diff report (dist/template-changes.md) for validation
"""

import os
import re
import shutil
from pathlib import Path
from docx import Document

# Define the casing mappings (original -> standardized)
# Only changes casing, preserves the actual text
CASING_MAPPINGS = {
    "{{DATE}}": "{{Date}}",
    "{{EXACT LEGAL NAME OF CLIENT}}": "{{Exact legal name of client}}",
    "{{FULL LEGAL NAME OF RECIPIENT}}": "{{Full legal name of Recipient}}",
}

# Templates to process (v1 scope)
TEMPLATES_V1 = [
    "Buyside Client Release Letter_EN_23-May-2025.docx",
    "Buyside Third Party Access Letter_EN_23-May-2025.docx",
    "Sellside Factual Accuracy Letter_EN_28-Oct-2024.docx",
]


def extract_all_text(doc: Document) -> list[tuple[str, str, str]]:
    """Extract all text from document with location info.

    Returns list of (location, paragraph_index, text) tuples.
    """
    texts = []

    # Body paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            texts.append(("body", f"para_{i}", para.text))

    # Tables
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    if para.text.strip():
                        texts.append(("table", f"table_{t_idx}_row_{r_idx}_cell_{c_idx}_para_{p_idx}", para.text))

    # Headers
    for s_idx, section in enumerate(doc.sections):
        for p_idx, para in enumerate(section.header.paragraphs):
            if para.text.strip():
                texts.append(("header", f"section_{s_idx}_header_para_{p_idx}", para.text))
        for p_idx, para in enumerate(section.footer.paragraphs):
            if para.text.strip():
                texts.append(("footer", f"section_{s_idx}_footer_para_{p_idx}", para.text))

    return texts


def replace_in_runs(para, old_text: str, new_text: str) -> bool:
    """Replace text in paragraph runs, handling split placeholders.

    Returns True if replacement was made.
    """
    full_text = para.text
    if old_text not in full_text:
        return False

    # Strategy: rebuild the paragraph text with replacement
    new_full_text = full_text.replace(old_text, new_text)

    # Find which runs contain parts of the old text
    # and distribute the new text across runs
    runs = list(para.runs)
    if not runs:
        return False

    # Simple approach: put all text in first run, clear others
    # This preserves the first run's formatting
    if len(runs) == 1:
        runs[0].text = new_full_text
    else:
        # Check if the placeholder is entirely within one run
        for run in runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                return True

        # Placeholder is split across runs - need to coalesce
        # Put all text in first run, clear others
        runs[0].text = new_full_text
        for run in runs[1:]:
            run.text = ""

    return True


def process_paragraphs(paragraphs, changes: list) -> None:
    """Process a list of paragraphs, applying casing mappings."""
    for para in paragraphs:
        for old_text, new_text in CASING_MAPPINGS.items():
            if old_text in para.text:
                original = para.text
                if replace_in_runs(para, old_text, new_text):
                    changes.append({
                        "original": old_text,
                        "new": new_text,
                        "context_before": original[:100],
                        "context_after": para.text[:100],
                    })


def process_template(input_path: Path, output_path: Path) -> list[dict]:
    """Process a single template, return list of changes made."""
    doc = Document(input_path)
    changes = []

    # Body paragraphs
    process_paragraphs(doc.paragraphs, changes)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, changes)

    # Headers and footers
    for section in doc.sections:
        process_paragraphs(section.header.paragraphs, changes)
        process_paragraphs(section.footer.paragraphs, changes)

    # Save processed document
    doc.save(output_path)

    return changes


def generate_diff_report(all_changes: dict[str, list[dict]], output_path: Path) -> None:
    """Generate markdown diff report."""
    lines = ["# Template Changes Report\n"]
    lines.append("This report shows all placeholder casing changes made during preprocessing.\n")
    lines.append("**Review carefully** to verify ONLY casing was changed (no content alterations).\n")
    lines.append("---\n")

    for template_name, changes in all_changes.items():
        lines.append(f"\n## {template_name}\n")

        if not changes:
            lines.append("*No changes required - all placeholders already have correct casing.*\n")
            continue

        lines.append(f"\n### Changes ({len(changes)} replacement{'s' if len(changes) != 1 else ''})\n")
        lines.append("| # | Original | Changed To |")
        lines.append("|---|----------|------------|")

        for i, change in enumerate(changes, 1):
            lines.append(f"| {i} | `{change['original']}` | `{change['new']}` |")

        lines.append("\n### Context Diffs\n")
        lines.append("```diff")
        for change in changes:
            lines.append(f"- {change['context_before']}...")
            lines.append(f"+ {change['context_after']}...")
        lines.append("```\n")

    lines.append("\n---\n")
    lines.append("## Summary\n")
    total = sum(len(c) for c in all_changes.values())
    lines.append(f"- **Total changes:** {total}")
    lines.append(f"- **Templates processed:** {len(all_changes)}")
    lines.append("\n**Validation:** If any changes appear to modify content (not just casing), do not proceed.\n")

    output_path.write_text("\n".join(lines))


def main():
    # Paths
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    templates_dir = project_root / "templates"
    dist_dir = project_root / "dist"

    # Ensure dist directory exists
    dist_dir.mkdir(exist_ok=True)

    print("=" * 60)
    print("Template Preprocessing Script")
    print("=" * 60)
    print(f"\nSource: {templates_dir}")
    print(f"Output: {dist_dir}")
    print(f"\nCasing mappings:")
    for old, new in CASING_MAPPINGS.items():
        print(f"  {old} -> {new}")
    print()

    all_changes = {}

    for template_name in TEMPLATES_V1:
        input_path = templates_dir / template_name
        output_path = dist_dir / template_name

        if not input_path.exists():
            print(f"WARNING: Template not found: {input_path}")
            continue

        print(f"Processing: {template_name}")
        changes = process_template(input_path, output_path)
        all_changes[template_name] = changes
        print(f"  -> {len(changes)} change(s) made")
        print(f"  -> Saved to: {output_path}")

    # Generate diff report
    diff_report_path = dist_dir / "template-changes.md"
    generate_diff_report(all_changes, diff_report_path)
    print(f"\nDiff report: {diff_report_path}")

    print("\n" + "=" * 60)
    print("DONE - Review dist/template-changes.md before proceeding")
    print("=" * 60)


if __name__ == "__main__":
    main()
