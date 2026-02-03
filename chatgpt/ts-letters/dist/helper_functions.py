"""
Helper Functions for DD Template Letter Drafter

This file provides functions for:
1. Replacing placeholders in DOCX templates
2. Validating no placeholders remain
3. Document text extraction for verification
"""

import re
from docx import Document


def replace_placeholders(doc_path: str, variables: dict[str, str]) -> Document:
    """
    Replace {{KEY}} placeholders throughout a DOCX document.

    Args:
        doc_path: Path to the DOCX template file
        variables: Dictionary mapping placeholder keys to replacement values
                   e.g., {"{{Date}}": "January 15, 2026", "{{Client Name}}": "Acme Corp"}

    Returns:
        Modified Document object (call .save() to persist)
    """
    doc = Document(doc_path)

    def replace_in_paragraph(para):
        """Replace placeholders in a single paragraph."""
        for key, value in variables.items():
            if key in para.text:
                # Replace in each run to preserve formatting
                for run in para.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Process body paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    # Process headers and footers
    for section in doc.sections:
        # Header
        for para in section.header.paragraphs:
            replace_in_paragraph(para)
        # Footer
        for para in section.footer.paragraphs:
            replace_in_paragraph(para)

    return doc


def validate_no_placeholders(doc: Document) -> list[str]:
    """
    Check document for any remaining {{...}} placeholders.

    Args:
        doc: Document object to check

    Returns:
        List of remaining placeholder patterns found (empty if clean)
    """
    remaining = []
    pattern = re.compile(r'\{\{[^}]+\}\}')

    def scan_paragraph(para):
        matches = pattern.findall(para.text)
        remaining.extend(matches)

    # Scan body
    for para in doc.paragraphs:
        scan_paragraph(para)

    # Scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    scan_paragraph(para)

    # Scan headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            scan_paragraph(para)
        for para in section.footer.paragraphs:
            scan_paragraph(para)

    # Return unique placeholders
    return list(set(remaining))


def extract_all_text(doc: Document) -> str:
    """
    Extract all text content from a document for verification.

    Args:
        doc: Document object

    Returns:
        Concatenated text from all paragraphs, tables, headers, footers
    """
    texts = []

    # Body
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        texts.append(para.text)

    # Headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            if para.text.strip():
                texts.append(f"[HEADER] {para.text}")
        for para in section.footer.paragraphs:
            if para.text.strip():
                texts.append(f"[FOOTER] {para.text}")

    return "\n".join(texts)


def generate_document(template_path: str, output_path: str, variables: dict[str, str]) -> tuple[bool, str]:
    """
    Generate a filled document from a template.

    Args:
        template_path: Path to the template DOCX
        output_path: Path to save the filled document
        variables: Dictionary of placeholder -> value mappings

    Returns:
        Tuple of (success: bool, message: str)
        - If success: (True, "Document generated successfully")
        - If failure: (False, "Error message with remaining placeholders")
    """
    try:
        # Load and fill template
        doc = replace_placeholders(template_path, variables)

        # Validate no placeholders remain
        remaining = validate_no_placeholders(doc)
        if remaining:
            return (False, f"ERROR: The following placeholders were not replaced: {remaining}")

        # Save document
        doc.save(output_path)
        return (True, f"Document generated successfully: {output_path}")

    except Exception as e:
        return (False, f"ERROR: Failed to generate document: {str(e)}")
